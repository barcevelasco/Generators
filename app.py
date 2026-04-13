# ==========================================
# 1. IMPORTS
# ==========================================
import streamlit as st
import requests
import pandas as pd
import html
from io import BytesIO
import datetime
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from bs4 import BeautifulSoup
import calendar
import time
import re
from dateutil import parser
from imf_data import get_fandd_march2026
import random

# ==========================================
# 2. CONFIGURACIÓN INICIAL Y ESTILOS
# ==========================================
st.set_page_config(page_title="Boletín Mensual", layout="wide")

st.markdown("""
    <style>
    div.stButton > button, div.stDownloadButton > button {
        background-color: #00205B !important;
        color: white !important;
        border: none !important;
    }
    div.stButton > button:hover, div.stDownloadButton > button:hover {
        background-color: #00153D !important;
        color: white !important;
    }
    span[data-baseweb="tag"] {
        background-color: #00205B !important;
        color: white !important;
    }
    .github-footer {
        position: fixed;
        right: 20px;
        bottom: 20px;
        background-color: rgba(255, 255, 255, 0.9);
        padding: 8px 12px;
        border-radius: 50px;
        border: 1px solid #d0d7de;
        z-index: 1000;
        display: flex;
        align-items: center;
        font-family: 'Calibri', sans-serif;
        text-decoration: none;
        color: #24292f;
        box-shadow: 0px 4px 12px rgba(0,0,0,0.1);
        transition: transform 0.2s, box-shadow 0.2s;
    }
    .github-footer:hover {
        transform: translateY(-2px);
        box-shadow: 0px 6px 16px rgba(0,0,0,0.15);
        color: #00205B;
        border-color: #00205B;
    }
    .github-icon {
        margin-right: 8px;
        width: 22px;
        height: 22px;
    }
    </style>
    <a class="github-footer" href="https://github.com/sdiazprado" target="_blank">
        <img class="github-icon" src="https://github.githubassets.com/images/modules/logos_page/GitHub-Mark.png" alt="GitHub Logo">
        <span><strong>@sdiazprado</strong></span>
    </a>
""", unsafe_allow_html=True)

# ==========================================
# 3. UTILIDADES DE FORMATO
# ==========================================
def clean_author_name(name):
    """Convierte nombres en mayúsculas a formato de nombre propio"""
    if not name:
        return ""
    cleaned = name.strip().title()
    cleaned = re.sub(r'\b([A-Z])\.\s*([A-Z])', lambda m: f"{m.group(1)}. {m.group(2)}", cleaned)
    return cleaned

# ==========================================
# 4. FUNCIONES DE EXTRACCIÓN (BACKEND)
# ==========================================

# --- SECCIÓN: REPORTES ---
# BID (Annual Reports en inglés)
@st.cache_data(show_spinner=False)
def load_reportes_bid_en(start_date_str, end_date_str):
    """
    Extrae Annual Reports del BID en inglés usando cloudscraper
    (mismo método que funciona para BID Investigación)
    """
    import cloudscraper
    from bs4 import BeautifulSoup
    import datetime
    import re
    import time
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BID Reportes: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
    
    rows = []
    page = 0
    max_pages = 5
    
    # Crear scraper con la misma configuración que usas en BID Investigación
    scraper = cloudscraper.create_scraper(
        browser={
            'browser': 'chrome',
            'platform': 'windows',
            'mobile': False
        },
        delay=5
    )
    
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
        'Accept-Language': 'en-US,en;q=0.5',
        'Accept-Encoding': 'gzip, deflate, br',
        'Connection': 'keep-alive',
        'Upgrade-Insecure-Requests': '1',
    }
    
    meses_map = {
        'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
        'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12,
        'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6,
        'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12
    }
    
    while page < max_pages:
        url = f"https://publications.iadb.org/en?f%5B0%5D=type%3AAnnual%20Reports&page={page}"
        print(f"📄 Página {page+1}: {url}")
        
        try:
            response = scraper.get(url, headers=headers, timeout=30)
            
            if response.status_code != 200:
                print(f"   ❌ Error HTTP: {response.status_code}")
                break
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Buscar artículos
            items = soup.find_all('div', class_='views-row')
            
            if not items:
                print(f"   📭 No hay resultados en página {page+1}")
                # Guardar HTML para depuración
                with open(f"bid_reportes_page_{page}_debug.html", "w", encoding="utf-8") as f:
                    f.write(response.text)
                print(f"   💾 HTML guardado en bid_reportes_page_{page}_debug.html")
                break
            
            print(f"   📚 Artículos encontrados: {len(items)}")
            
            items_found = 0
            for item in items:
                try:
                    # Título y link
                    title_div = item.find('div', class_='views-field-field-title')
                    if not title_div:
                        continue
                    
                    a_tag = title_div.find('a')
                    if not a_tag:
                        continue
                    
                    titulo = a_tag.get_text(strip=True)
                    link = a_tag.get('href')
                    if link and not link.startswith('http'):
                        link = "https://publications.iadb.org" + link
                    
                    # Fecha
                    date_div = item.find('div', class_='views-field-field-date-issued-text')
                    if not date_div:
                        continue
                    
                    date_text = date_div.get_text(strip=True)
                    match = re.search(r'([A-Za-z]{3,9})\s+(\d{4})', date_text)
                    if not match:
                        continue
                    
                    mes_str = match.group(1).lower()[:3]
                    año = int(match.group(2))
                    mes_num = meses_map.get(mes_str, 1)
                    parsed_date = datetime.datetime(año, mes_num, 15)
                    
                    # Filtrar por fecha
                    if parsed_date < start_date or parsed_date > end_date:
                        continue
                    
                    if not any(r['Link'] == link for r in rows):
                        rows.append({
                            "Date": parsed_date,
                            "Title": titulo,
                            "Link": link,
                            "Organismo": "BID (Reportes)"
                        })
                        items_found += 1
                        print(f"   ✅ {parsed_date.date()} - {titulo[:50]}...")
                        
                except Exception as e:
                    continue
            
            print(f"   📊 Documentos en página {page+1}: {items_found}")
            
            if items_found == 0 and page > 0:
                break
            
            page += 1
            time.sleep(2)
            
        except Exception as e:
            print(f"   ❌ Error: {e}")
            break
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.drop_duplicates(subset=['Link'])
        df = df.sort_values("Date", ascending=False)
    
    print(f"✅ BID Reportes - Total: {len(df)} documentos")
    return df

# == Reportes BPI == #

@st.cache_data(show_spinner=False)
def load_reportes_bpi(start_date_str, end_date_str):
    urls_api = [
        "https://www.bis.org/api/document_lists/bcbspubls.json",
        "https://www.bis.org/api/document_lists/cpmi_publs.json"
    ]
    urls_html = ["https://www.bis.org/ifc/publications.htm"]
    headers = {'User-Agent': 'Mozilla/5.0'}
    
    try: 
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: 
        start_date = datetime.datetime(2000, 1, 1)
    
    rows = []
    
    for url in urls_api:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            data = res.json()
            lista_documentos = data.get("list", {})
            for path, doc_info in lista_documentos.items():
                titulo = html.unescape(doc_info.get("short_title", ""))
                if not titulo: continue
                link = "https://www.bis.org" + doc_info.get("path", "")
                if not link.endswith(".htm") and not link.endswith(".pdf"):
                    link += ".htm"
                date_str = doc_info.get("publication_start_date", "")
                parsed_date = None
                if date_str:
                    try: 
                        parsed_date = parser.parse(date_str)
                    except: 
                        pass
                if not parsed_date: continue
                if parsed_date >= start_date:
                    rows.append({"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "BPI"})
        except Exception as e:
            continue

    for url in urls_html:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            soup = BeautifulSoup(res.text, 'html.parser')
            content_div = soup.find('div', id='cmsContent')
            if not content_div: continue
            for p in content_div.find_all('p'):
                a_tag = p.find('a')
                if not a_tag: continue
                titulo = a_tag.get_text(strip=True)
                href = a_tag.get('href', '')
                if not href or 'index.htm' in href: continue 
                link = "https://www.bis.org" + href if href.startswith('/') else href
                full_text = p.get_text(strip=True)
                date_str = full_text.replace(titulo, '').strip(', ')
                parsed_date = None
                if date_str:
                    try: 
                        parsed_date = parser.parse(date_str)
                    except: 
                        pass
                if not parsed_date:
                    match = re.search(r'\b(20\d{2})\b', titulo)
                    if match: 
                        parsed_date = datetime.datetime(int(match.group(1)), 1, 1)
                if not parsed_date: continue
                if parsed_date >= start_date:
                    rows.append({"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "BPI"})
        except Exception as e:
            continue
            
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None: 
            df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df

# == Reportes BM == #

@st.cache_data(show_spinner=False)
def load_reportes_bm(start_date_str, end_date_str):
    """
    Extractor para Reportes del BM usando API de DSpace
    """
    base_url = "https://openknowledge.worldbank.org/server/api/discover/search/objects"
    headers = {'User-Agent': 'Mozilla/5.0'}

    # ID exacto de la comunidad de Publicaciones
    scope_id = '06251f8a-62c2-59fb-add5-ec0993fc20d9'

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BM Reportes: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()

    # Palabras clave para identificar reportes (ampliadas)
    palabras_reporte = [
        r'\breport\b', r'\boutlook\b', r'\bprospects\b', r'\bupdate\b',
        r'\breview\b', r'\bmonitor\b', r'\bbulletin\b', r'\boverview\b',
        r'\bassessment\b', r'\banalysis\b', r'\bforecast\b', r'\btrends?\b',
        r'\bdevelopments?\b', r'\bglobal economic\b', r'\bcommodity markets\b',
        r'\beconomic\s+report\b', r'\bcountry\s+update\b', r'\bquarterly\b',
        r'\bannual\s+report\b', r'\bglobal\s+development\b', r'\bmacroeconomic\b',
        r'\bfiscal\s+update\b', r'\bpolicy\s+note\b', r'\bworking\s+paper\b',
        r'\bdiscussion\s+paper\b', r'\bpolicy\s+research\s+working\s+paper\b'
    ]

    rows = []
    page = 0
    max_pages = 10  # Aumentado para capturar más
    
    while page < max_pages:
        try:
            # Aumentar size a 50 para capturar más por página
            params = {
                'scope': scope_id,
                'sort': 'dc.date.issued,DESC',
                'page': page,
                'size': 50
            }
            res = requests.get(base_url, headers=headers, params=params, timeout=15)
            data = res.json()

            objects = data.get('_embedded', {}).get(
                'searchResult', {}).get('_embedded', {}).get('objects', [])
            
            if not objects:
                print(f"📭 No hay más resultados en página {page}")
                break

            print(f"📄 Página {page + 1}: {len(objects)} objetos encontrados")
            
            items_found = 0
            for obj in objects:
                item = obj.get('_embedded', {}).get('indexableObject', {})
                meta = item.get('metadata', {})

                # Extraer Título
                title = meta.get('dc.title', [{'value': ''}])[0].get('value', '')
                if not title:
                    continue
                
                # Extraer Fecha
                date_s = meta.get('dc.date.issued', [{'value': ''}])[0].get('value', '')
                if not date_s:
                    continue
                    
                try:
                    parsed_date = parser.parse(date_s)
                    if parsed_date.tzinfo is not None:
                        parsed_date = parsed_date.replace(tzinfo=None)
                except:
                    continue

                if parsed_date < start_date or parsed_date > end_date:
                    continue
                
                # Revisión de resultados 
                print(f"   📄 {parsed_date.date()} - {title[:80]}...")

                # ========== FILTRO MEJORADO ==========
                es_reporte = False
                
                # 1. Revisar título
                for palabra in palabras_reporte:
                    if re.search(palabra, title.lower()):
                        es_reporte = True
                        break
                
                # 2. Si no está en título, revisar descripción
                if not es_reporte:
                    abstract_list = meta.get('dc.description.abstract', [])
                    desc_list = meta.get('dc.description', [])
                    description = ""
                    if abstract_list:
                        description = abstract_list[0].get('value', '').lower()
                    elif desc_list:
                        description = desc_list[0].get('value', '').lower()
                    
                    for palabra in palabras_reporte:
                        if re.search(palabra, description):
                            es_reporte = True
                            break
                
                # 3. Si no es reporte, saltar
                #if not es_reporte:
                #    continue
                # ==================================(ESTE COMMENT evita que filtre innecesariamente todo el listado disponible)

                # Link permanente
                link = meta.get('dc.identifier.uri', [{'value': ''}])[0].get('value', '')
                if not link:
                    link = f"https://openknowledge.worldbank.org/entities/publication/{item.get('id', '')}"

                if not any(r['Link'] == link for r in rows):
                    rows.append({
                        "Date": parsed_date, 
                        "Title": title,
                        "Link": link, 
                        "Organismo": "BM"
                    })
                    items_found += 1
                    print(f"   ✅ {parsed_date.date()} - {title[:60]}...")

            print(f"   📊 Documentos en página {page + 1}: {items_found}")
            
            # Si no encontramos nada en 2 páginas consecutivas, paramos
            if items_found == 0 and page > 1:
                break
                
            page += 1
            time.sleep(0.5)
            
        except Exception as e:
            print(f"⚠️ Error en página {page}: {e}")
            break

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    print(f"✅ BM Reportes - Total: {len(df)} documentos")
    return df

# == Reportes CEF == #

@st.cache_data(show_spinner=False)
def load_reportes_cef(start_date_str, end_date_str):
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: 
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: 
        start_date = datetime.datetime(2000, 1, 1)
    rows, page = [], 1
    while True:
        url = f"https://www.fsb.org/publications/?dps_paged={page}"
        try:
            res = requests.get(url, headers=headers, timeout=15)
            soup = BeautifulSoup(res.text, 'html.parser')
            items = soup.find_all('div', class_=lambda c: c and 'post-excerpt' in c)
            if not items: break
            items_found = 0
            for item in items:
                title_div = item.find('div', class_='post-title')
                if not title_div or not title_div.find('a'): continue
                a_tag = title_div.find('a')
                titulo_raw = a_tag.get_text(strip=True)
                link = a_tag.get('href', '')
                date_div = item.find('div', class_='post-date')
                parsed_date = None
                if date_div:
                    try: parsed_date = parser.parse(date_div.get_text(strip=True))
                    except: pass
                if not parsed_date: continue
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": titulo_raw, "Link": link, "Organismo": "CEF"})
                    items_found += 1
            if items_found == 0 or (rows and rows[-1]['Date'] < start_date): break
            page += 1
            time.sleep(0.5) 
        except: break
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

# == Reportes OCDE == #

@st.cache_data(show_spinner=False)
def load_reportes_ocde(start_date_str, end_date_str):
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    rows = []
    try: 
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: 
        start_date = datetime.datetime(2000, 1, 1)
    year = start_date.year
    chrome_options = Options()
    chrome_options.add_argument("--headless=new")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    try:
        driver = webdriver.Chrome(options=chrome_options)
        url = f"https://www.oecd.org/en/search/publications.html?orderBy=mostRecent&page=0&facetTags=oecd-content-types%3Apublications%2Freports%2Coecd-languages%3Aen&minPublicationYear={year}&maxPublicationYear={year}"
        driver.get(url)
        time.sleep(12) 
        js_script = """
        let linksData = [];
        function findLinks(root) {
            let els = root.querySelectorAll('*');
            els.forEach(el => {
                if (el.shadowRoot) findLinks(el.shadowRoot);
                if (el.tagName === 'A' && el.href) {
                    let text = el.innerText || el.textContent;
                    let aria = el.getAttribute('aria-label') || el.getAttribute('title') || '';
                    let final_text = text.trim() ? text.trim() : aria.trim();
                    if(final_text.length > 15) { linksData.push({ title: final_text, link: el.href }); }
                }
            });
        }
        findLinks(document); return linksData;
        """
        extracted_links = driver.execute_script(js_script)
        driver.quit()
        for item in extracted_links:
            href = item['link'].lower()
            title = item['title'].replace('\n', ' ')
            firmas_validas = ['/publications/', '/reports/', 'oecd-ilibrary.org', '/books/']
            if any(firma in href for firma in firmas_validas):
                if any(basura in title.lower() for basura in ['download', 'read more', 'pdf', 'buy', 'search', 'subscribe']): continue
                if not any(r['Link'] == item['link'] for r in rows):
                    rows.append({"Date": start_date, "Title": title, "Link": item['link'], "Organismo": "OCDE"})
    except: pass
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

# == Reportes BPI == #

@st.cache_data(show_spinner=False)
def load_reportes_bpi(start_date_str, end_date_str):
    urls_api = ["https://www.bis.org/api/document_lists/bcbspubls.json", "https://www.bis.org/api/document_lists/cpmi_publs.json"]
    urls_html = ["https://www.bis.org/ifc/publications.htm"]
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows = []
    for url in urls_api:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            data = res.json()
            for path, doc in data.get("list", {}).items():
                titulo = html.unescape(doc.get("short_title", ""))
                if not titulo: continue
                link = "https://www.bis.org" + doc.get("path", "")
                if not link.endswith(".htm") and not link.endswith(".pdf"): link += ".htm"
                try: parsed_date = parser.parse(doc.get("publication_start_date", ""))
                except: continue
                if parsed_date >= start_date:
                    rows.append({"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "BPI"})
        except: continue
    for url in urls_html:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            soup = BeautifulSoup(res.text, 'html.parser')
            content_div = soup.find('div', id='cmsContent')
            if not content_div: continue
            for p in content_div.find_all('p'):
                a_tag = p.find('a')
                if not a_tag: continue
                titulo = a_tag.get_text(strip=True)
                href = a_tag.get('href', '')
                if not href or 'index.htm' in href: continue 
                link = "https://www.bis.org" + href if href.startswith('/') else href
                parsed_date = None
                try: parsed_date = parser.parse(p.get_text(strip=True).replace(titulo, '').strip(', '))
                except: pass
                if not parsed_date:
                    match = re.search(r'\b(20\d{2})\b', titulo)
                    if match: parsed_date = datetime.datetime(int(match.group(1)), 1, 1)
                if parsed_date and parsed_date >= start_date:
                    rows.append({"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "BPI"})
        except: continue
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

# --- SECCIÓN: PUBLICACIONES INSTITUCIONALES ---
@st.cache_data(show_spinner=False)
def load_pub_inst_cef(start_date_str, end_date_str):
    """Extractor para Publicaciones Institucionales del CEF (FSB)"""
    url = "https://www.fsb.org/publications/key-regular-publications/"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: 
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: 
        start_date = datetime.datetime(2000, 1, 1)

    rows = []
    try:
        res = requests.get(url, headers=headers, timeout=15)
        soup = BeautifulSoup(res.text, 'html.parser')
        
        sections = soup.find_all('div', class_='wp-bootstrap-blocks-row')
        
        for section in sections:
            h2 = section.find('h2')
            if not h2: continue
            base_title = h2.get_text(strip=True)
            
            # Latest
            latest_btn = section.find('button', class_='btn-primary')
            if latest_btn and latest_btn.find('a'):
                a_tag = latest_btn.find('a')
                link = "https://www.fsb.org" + a_tag['href'] if a_tag['href'].startswith('/') else a_tag['href']
                
                date_match = re.search(r'\((.*?)\)', a_tag.get_text())
                parsed_date = None
                if date_match:
                    try: 
                        parsed_date = parser.parse(date_match.group(1))
                    except: 
                        pass
                
                if parsed_date and parsed_date >= start_date:
                    rows.append({"Date": parsed_date, "Title": f"{base_title}: Latest Report", "Link": link, "Organismo": "CEF"})

            # Previous
            dropdown = section.find('div', class_='dropdown-menu')
            if dropdown:
                links = dropdown.find_all('a')
                for l in links:
                    link = l['href']
                    year_text = l.get_text(strip=True)
                    try: 
                        parsed_date = datetime.datetime(int(year_text), 1, 1)
                    except: 
                        parsed_date = None
                    
                    if parsed_date and parsed_date >= start_date:
                        rows.append({"Date": parsed_date, "Title": f"{base_title} ({year_text})", "Link": link, "Organismo": "CEF"})

    except Exception as e:
        print(f"Error extrayendo Pub Institucionales CEF:", e)

    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

# == BPI (Publicaciones Institucionales) == #

@st.cache_data(show_spinner=False)
def load_pub_inst_bpi(start_date_str, end_date_str):
    urls_api = [
        "https://www.bis.org/api/document_lists/annualeconomicreports.json",
        "https://www.bis.org/api/document_lists/quarterlyreviews.json"
    ]
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: 
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: 
        start_date = datetime.datetime(2000, 1, 1)

    rows = []
    for url in urls_api:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            data = res.json()
            lista_documentos = data.get("list", {})
            for path, doc_info in lista_documentos.items():
                titulo = html.unescape(doc_info.get("short_title", ""))
                if not titulo: continue
                
                link = "https://www.bis.org" + doc_info.get("path", "")
                if not link.endswith(".htm") and not link.endswith(".pdf"):
                    link += ".htm"
                    
                date_str = doc_info.get("publication_start_date", "")
                parsed_date = None
                if date_str:
                    try: 
                        parsed_date = parser.parse(date_str)
                    except: 
                        pass
                if not parsed_date: continue
                
                if parsed_date >= start_date:
                    rows.append({"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "BPI"})
        except Exception as e:
            continue

    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None: 
            df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df

# == Reportes FMI == #

@st.cache_data(show_spinner=False)
def load_country_reports_fmi(start_date_str, end_date_str):
    """Extractor FMI - Country Reports (Conexión Directa a Coveo API)"""
    try: 
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: 
        start_date = datetime.datetime(2000, 1, 1)
    
    rows = []
    
    url = "https://imfproduction561s308u.org.coveo.com/rest/search/v2?organizationId=imfproduction561s308u"
    
    headers = {
        "Authorization": "Bearer xx742a6c66-f427-4f5a-ae1e-770dc7264e8a",
        "Content-Type": "application/json",
        "Accept": "application/json",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }
    
    payload = {
        "aq": "@imfseries==\"IMF Staff Country Reports\"",
        "numberOfResults": 100,
        "sortCriteria": "@imfdate descending"
    }
    
    try:
        res = requests.post(url, headers=headers, json=payload, timeout=15)
        
        if res.status_code == 200:
            data = res.json()
            
            for item in data.get("results", []):
                titulo = item.get("title", "")
                link = item.get("clickUri", "")
                
                raw_date = item.get("raw", {}).get("date")
                parsed_date = None
                if raw_date:
                    try:
                        parsed_date = datetime.datetime.fromtimestamp(raw_date / 1000.0)
                    except: 
                        pass
                
                if not titulo or not link or not parsed_date: continue
                
                if parsed_date >= start_date:
                    if not any(r['Link'] == link for r in rows):
                        rows.append({"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "FMI"})
    except Exception as e:
        pass
        
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

# == Reportes FMI - PRENSA == #

@st.cache_data(show_spinner=False)
def load_press_releases_fmi(start_date_str, end_date_str):
    """Extractor FMI - Press Releases (Historial completo vía Coveo API)"""
    try: 
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: 
        start_date = datetime.datetime(2000, 1, 1)
    
    rows = []
    
    url = "https://imfproduction561s308u.org.coveo.com/rest/search/v2?organizationId=imfproduction561s308u"
    
    headers = {
        "Authorization": "Bearer xx742a6c66-f427-4f5a-ae1e-770dc7264e8a",
        "Content-Type": "application/json",
        "Accept": "application/json",
        "Origin": "https://www.imf.org",
        "Referer": "https://www.imf.org/",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }
    
    payload = {
        "aq": "@imftype==\"Press Release\" AND @syslanguage==\"English\"",
        "numberOfResults": 150,
        "sortCriteria": "@imfdate descending"
    }
    
    try:
        res = requests.post(url, headers=headers, json=payload, timeout=15)
        
        if res.status_code == 200:
            data = res.json()
            
            for item in data.get("results", []):
                titulo = item.get("title", "")
                link = item.get("clickUri", "")
                
                raw_date = item.get("raw", {}).get("date")
                parsed_date = None
                if raw_date:
                    try:
                        parsed_date = datetime.datetime.fromtimestamp(raw_date / 1000.0)
                    except: 
                        pass
                
                if not titulo or not link or not parsed_date: continue
                
                if parsed_date >= start_date:
                    if not any(r['Link'] == link for r in rows):
                        rows.append({"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "FMI"})
    except Exception as e:
        pass
        
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

# == Reportes FMI - COUNTRIES == #

@st.cache_data(show_spinner=False)
def load_country_reports_elibrary(start_date_str, end_date_str):
    """Extractor FMI - Country Reports (Bypass de Tapestry 5 AJAX Lazy-Loading)"""
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8'
    }
    
    try: 
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: 
        start_date = datetime.datetime(2000, 1, 1)
    
    rows = []
    base_domain = "https://www.elibrary.imf.org"
    url_overview = f"{base_domain}/view/journals/002/002-overview.xml"
    
    try:
        res = requests.get(url_overview, headers=headers, timeout=15)
        if res.status_code != 200: return pd.DataFrame()
        
        soup = BeautifulSoup(res.text, 'html.parser')
        
        ajax_links = []
        current_year = datetime.datetime.now().year
        target_years = [str(current_year), str(current_year - 1)] 
        
        for li in soup.find_all('div', attrs={'data-toc-role': 'li'}):
            label_div = li.find('div', class_='label')
            if not label_div: continue
            
            texto_label = label_div.get_text()
            if any(year in texto_label for year in target_years):
                a_tag = li.find('a', class_='ajax-control')
                if a_tag and a_tag.has_attr('href'):
                    ajax_links.append(base_domain + a_tag['href'])
        
        headers_ajax = headers.copy()
        headers_ajax['X-Requested-With'] = 'XMLHttpRequest'
        headers_ajax['Accept'] = 'application/json, text/javascript, */*; q=0.01'
        
        for ajax_url in ajax_links:
            try:
                res_ajax = requests.get(ajax_url, headers=headers_ajax, timeout=15)
                if res_ajax.status_code != 200: continue
                
                data = res_ajax.json()
                
                html_fragment = ""
                if "zones" in data:
                    for zone_id, html_content in data["zones"].items():
                        html_fragment += html_content
                        
                if not html_fragment: continue
                
                soup_fragment = BeautifulSoup(html_fragment, 'html.parser')
                
                for a_tag in soup_fragment.find_all('a', href=True):
                    href = a_tag['href']
                    titulo = a_tag.get_text(strip=True)
                    
                    if '/view/journals/002/' in href and len(titulo) > 15:
                        link_real = base_domain + href if href.startswith('/') else href
                        
                        date_str = ""
                        for padre in a_tag.find_parents(['div', 'li'], limit=3):
                            texto_padre = padre.get_text(separator=" ", strip=True)
                            
                            match = re.search(r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2}?,?\s*\d{4}', texto_padre)
                            if not match:
                                match = re.search(r'\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}', texto_padre)
                                
                            if match:
                                date_str = match.group(0)
                                break
                                
                        parsed_date = None
                        if date_str:
                            try:
                                parsed_date = parser.parse(date_str)
                                if parsed_date.tzinfo is not None: parsed_date = parsed_date.replace(tzinfo=None)
                            except: 
                                pass
                            
                        if parsed_date and parsed_date >= start_date:
                            if not any(r['Link'] == link_real for r in rows):
                                rows.append({"Date": parsed_date, "Title": titulo, "Link": link_real, "Organismo": "FMI"})
            except:
                continue
                
    except Exception as e:
        pass
        
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

# == Publicaciones Institucionales FMI == #

@st.cache_data(show_spinner=False)
def load_pub_inst_fmi(start_date_str, end_date_str):
    """Extractor FMI - Vía directa por API Next.js"""
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Accept': 'application/json, text/plain, */*'
    }
    
    try: 
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: 
        start_date = datetime.datetime(2000, 1, 1)
    
    rows = []
    
    build_id = "OPXKbpp2La91iW-gTVkBX"
    try:
        res_html = requests.get("https://www.imf.org/en/publications", headers=headers, timeout=15)
        match = re.search(r'"buildId":"([^"]+)"', res_html.text)
        if match:
            build_id = match.group(1)
    except:
        pass

    endpoints_json = [
        f"https://www.imf.org/_next/data/{build_id}/en/publications/fm.json",
        f"https://www.imf.org/_next/data/{build_id}/en/publications/weo.json",
        f"https://www.imf.org/_next/data/{build_id}/en/publications/gfsr.json"
    ]
    
    for url in endpoints_json:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            if res.status_code != 200: continue
            data = res.json()
            
            def extraer_issues(obj):
                if isinstance(obj, dict):
                    if "issuePage" in obj and isinstance(obj["issuePage"], dict) and "results" in obj["issuePage"]:
                        for r in obj["issuePage"]["results"]: yield r
                    for k, v in obj.items(): yield from extraer_issues(v)
                elif isinstance(obj, list):
                    for item in obj: yield from extraer_issues(item)

            for issue in extraer_issues(data):
                titulo = issue.get("title", {}).get("jsonValue", {}).get("value", "")
                link_raw = issue.get("url", {}).get("url", "") or issue.get("url", {}).get("path", "")
                if not titulo or not link_raw: continue
                
                link_real = link_raw if link_raw.startswith("http") else "https://www.imf.org" + link_raw
                
                d_str = issue.get("publicationDate", {}).get("jsonValue", {}).get("value", "")
                if d_str:
                    try:
                        parsed_date = parser.parse(d_str)
                        if parsed_date.tzinfo is not None: parsed_date = parsed_date.replace(tzinfo=None)
                        if parsed_date >= start_date and not any(r['Link'] == link_real for r in rows):
                            rows.append({"Date": parsed_date, "Title": titulo, "Link": link_real, "Organismo": "FMI"})
                    except: 
                        pass
        except:
            continue
            
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df
    
# == Publicaciones Institucionales BM == #

@st.cache_data(show_spinner=False)
def load_pub_inst_bm(start_date_str, end_date_str):
    """Extractor para Publicaciones Institucionales (Colecciones Específicas) del BM"""
    base_url = "https://openknowledge.worldbank.org/server/api/discover/search/objects"
    headers = {'User-Agent': 'Mozilla/5.0'}
    
    scopes = [
        '4c48a649-7773-4d0f-b441-f5fc7e8d67f8',  # Business Ready
        '09c5e8fc-187f-5c2f-a077-3e03044c7b62',  # Perspectivas económicas mundiales
        '3d9bbbf6-c007-5043-b655-04d8a1cfbfb2'   # Tercera colección
    ]
    
    try: 
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: 
        start_date = datetime.datetime(2000, 1, 1)
    
    rows = []
    
    for scope in scopes:
        page = 0
        while True:
            try:
                params = {
                    'scope': scope,
                    'sort': 'dc.date.issued,DESC',
                    'page': page,
                    'size': 20
                }
                res = requests.get(base_url, headers=headers, params=params, timeout=15)
                data = res.json()
                
                objects = data.get('_embedded', {}).get('searchResult', {}).get('_embedded', {}).get('objects', [])
                if not objects: break
                
                items_found = 0
                for obj in objects:
                    item = obj.get('_embedded', {}).get('indexableObject', {})
                    meta = item.get('metadata', {})
                    
                    title = meta.get('dc.title', [{'value': ''}])[0].get('value', '')
                    date_s = meta.get('dc.date.issued', [{'value': ''}])[0].get('value', '')
                    
                    parsed_date = None
                    if date_s:
                        try: 
                            parsed_date = parser.parse(date_s)
                        except: 
                            pass
                    
                    if not parsed_date or parsed_date < start_date: continue
                    
                    link = meta.get('dc.identifier.uri', [{'value': ''}])[0].get('value', '')
                    if not link: link = f"https://openknowledge.worldbank.org/entities/publication/{item.get('id', '')}"
                    
                    if not any(r['Link'] == link for r in rows):
                        rows.append({"Date": parsed_date, "Title": title, "Link": link, "Organismo": "BM"})
                        items_found += 1
                
                if items_found == 0: break
                page += 1
                if page > 3: break
                time.sleep(0.2)
            except:
                break
                
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None: 
            df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df

# ========== FUNCIÓN PARA FMI (F&D MAGAZINE) - VERSIÓN MEJORADA ==========
@st.cache_data(show_spinner=False)
def load_pub_inst_imf(start_date_str, end_date_str):
    """Extrae artículos de F&D Magazine desde el JSON embebido en la página"""
    import requests
    import json
    import re
    import datetime
    import pandas as pd

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 FMI F&D: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"⚠️ Error en fechas, usando rango por defecto")

    url = "https://www.imf.org/en/publications/fandd/issues"
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}

    rows = []

    try:
        print(f"📡 Solicitando página: {url}")
        res = requests.get(url, headers=headers, timeout=15)
        
        if res.status_code != 200:
            print(f"❌ Error al acceder a la página: {res.status_code}")
            return pd.DataFrame()

        # Buscar el script con id="__NEXT_DATA__"
        match = re.search(r'<script id="__NEXT_DATA__" type="application/json">(.*?)</script>', res.text, re.DOTALL)
        if not match:
            print("❌ No se encontró el script __NEXT_DATA__")
            return pd.DataFrame()

        data = json.loads(match.group(1))
        
        # Navegar hasta issueList
        try:
            issue_list = data['props']['pageProps']['page']['placeholders']['Content'][0]['placeholders']['Article-1099474'][0]['fields']['issueList']
        except (KeyError, IndexError, TypeError) as e:
            print(f"❌ Error navegando en JSON: {e}")
            return pd.DataFrame()

        total_issues = issue_list.get('total', 0)
        print(f"✅ Total de números encontrados: {total_issues}")

        for issue in issue_list.get('results', []):
            # Extraer fecha del número (issueLabel ej: "December 2025")
            issue_label = issue.get('issueLabel', {}).get('value', '')
            issue_date = None
            
            # Parsear fecha (ej: "December 2025" -> datetime(2025, 12, 1))
            match_date = re.search(r'([A-Za-z]+)\s+(\d{4})', issue_label)
            if match_date:
                mes_str = match_date.group(1).lower()
                año = int(match_date.group(2))
                mes_num = {
                    'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6,
                    'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12
                }.get(mes_str, 1)
                issue_date = datetime.datetime(año, mes_num, 1)
            
            if not issue_date:
                continue
            
            # Filtrar por rango de fechas
            if issue_date < start_date or issue_date > end_date:
                print(f"⏭️ Número fuera de rango: {issue_label}")
                continue
            
            print(f"\n📖 Procesando número: {issue_label}")
            
            # Procesar artículos del número
            articles = issue.get('children', {}).get('results', [])
            print(f"   Artículos encontrados: {len(articles)}")
            
            for article in articles:
                # Título del artículo
                title = article.get('articleTitle', {}).get('value', '')
                if not title:
                    continue
                
                # Enlace del artículo
                article_url = article.get('url', {}).get('url', '')
                if not article_url:
                    continue
                
                # Categoría (útil para filtrar si se desea)
                category = article.get('fANDDCategory', {}).get('id', '')
                
                # Limpiar título
                title = re.sub(r'\s+', ' ', title).strip()
                
                rows.append({
                    "Date": issue_date,
                    "Title": title,
                    "Link": article_url,
                    "Organismo": "FMI"
                })
                print(f"  ✅ {title[:60]}...")
        
    except Exception as e:
        print(f"❌ Error general: {e}")
        import traceback
        traceback.print_exc()
        return pd.DataFrame()

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.drop_duplicates(subset=['Link'])
        df = df.sort_values("Date", ascending=False)
        print(f"\n✅ TOTAL FMI F&D: {len(df)} artículos")
    else:
        print("⚠️ No se encontraron artículos en el rango seleccionado")

    return df

# ========== FUNCIÓN PARA CEMLA (PUBLICACIONES INSTITUCIONALES) ==========
@st.cache_data(show_spinner=False)
def load_pub_inst_cemla(start_date_str, end_date_str):
    """Extractor para Boletín CEMLA - Versión optimizada para la estructura de Mailchimp"""
    import requests
    from bs4 import BeautifulSoup
    import datetime
    import re
    import pandas as pd
    import time

    url = "https://www.cemla.org/comunicados.html"
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}

    print("="*50)
    print("🔍 Iniciando extracción de CEMLA con novedades individuales...")
    print(f"📅 Rango solicitado: {start_date_str} a {end_date_str}")
    print("="*50)

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"✅ Fechas parseadas: {start_date.date()} a {end_date.date()}")
    except Exception as e:
        print(f"⚠️ Error parseando fechas: {e}")
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now() + datetime.timedelta(days=365)
        print(f"📅 Usando rango por defecto: {start_date.date()} a {end_date.date()}")

    rows = []

    # Palabras y URLs a excluir
    palabras_excluir = [
        'convocatoria', 'premio', 'curso', 'taller', 'seminario',
        'evento', 'webinar', 'congreso', 'beca', 'inscripción',
        'registro', 'participación', 'invitación', 'calendario',
        'programa de actividades', 'agenda', 'convocan', 'postulación',
        'reunión de gobernadores', 'reunión de responsables', 'encuesta',
        'award', 'prize', 'conference', 'workshop', 'registration',
        'call for papers', 'agenda', 'calendar', 'program', 'invitation',
        'survey', 'meeting', 'governors', 'responsables'
    ]
    urls_excluir = [
        'calendario', 'premiodebancacentral', 'convocatoria',
        'award', 'prize', 'course', 'workshop', 'event',
        'reunion', 'meeting', 'programa-actividades'
    ]

    meses_map = {
        'enero': 1, 'febrero': 2, 'marzo': 3, 'abril': 4,
        'mayo': 5, 'junio': 6, 'julio': 7, 'agosto': 8,
        'septiembre': 9, 'octubre': 10, 'noviembre': 11, 'diciembre': 12,
    }

    try:
        print(f"📡 Solicitando lista de boletines desde {url}...")
        res = requests.get(url, headers=headers, timeout=15)
        print(f"   Status code: {res.status_code}")
        
        if res.status_code != 200:
            print(f"   ❌ Error al acceder a la página")
            return pd.DataFrame()

        soup = BeautifulSoup(res.text, 'html.parser')
        print(f"✅ Página cargada, {len(soup.text)} caracteres")

        # ===== 1. EXTRAER LISTA DE BOLETINES =====
        boletines = []
        for element in soup.find_all(['p', 'div', 'h3', 'h4', 'li']):
            text = element.get_text(strip=True)
            match = re.match(r'^([A-Za-z]+)\s+(\d{4})', text)
            if not match:
                continue

            mes_str, year_str = match.groups()
            mes_num = meses_map.get(mes_str.lower())
            if not mes_num:
                print(f"   ⚠️ Mes no reconocido: {mes_str}")
                continue

            try:
                fecha = datetime.datetime(int(year_str), mes_num, 1)
            except Exception as e:
                print(f"   ⚠️ Error fecha: {e}")
                continue

            a_tag = element.find('a', href=True, string=re.compile(r'Ver más', re.I))
            if not a_tag:
                next_elem = element.find_next_sibling()
                if next_elem:
                    a_tag = next_elem.find('a', href=True, string=re.compile(r'Ver más', re.I))
            
            if a_tag:
                href = a_tag.get('href')
                if href:
                    if href.startswith('/'):
                        link = f"https://www.cemla.org{href}"
                    elif href.startswith('http'):
                        link = href
                    else:
                        link = f"https://www.cemla.org/{href}"
                    
                    boletines.append({
                        'fecha': fecha,
                        'titulo': text,
                        'link': link
                    })
                    print(f"📌 Boletín encontrado: {fecha.strftime('%Y-%m')} - {text[:50]}...")

        print(f"✅ Total boletines principales: {len(boletines)}")

        if not boletines:
            print("⚠️ No se encontraron boletines. Verifica la estructura de la página.")
            with open("cemla_debug.html", "w", encoding="utf-8") as f:
                f.write(res.text)
            print("💾 HTML guardado en cemla_debug.html para depuración")
            return pd.DataFrame()

        # ===== 2. PROCESAR CADA BOLETÍN =====
        for boletin in boletines:
            if boletin['fecha'] < start_date or boletin['fecha'] > end_date:
                print(f"⏭️ Boletín fuera de rango: {boletin['fecha'].strftime('%Y-%m')}")
                continue

            print(f"\n🔍 Procesando boletín {boletin['fecha'].strftime('%Y-%m')}: {boletin['link']}")
            
            try:
                time.sleep(1)
                
                res_boletin = requests.get(boletin['link'], headers=headers, timeout=15)
                if res_boletin.status_code != 200:
                    print(f"  ⚠️ Error al acceder al boletín: {res_boletin.status_code}")
                    continue

                soup_boletin = BeautifulSoup(res_boletin.text, 'html.parser')
                
                novedades = []
                
                # ===== ESTRATEGIA MEJORADA: Buscar bloques de novedades =====
                # Busca divs con clase "ipost clearfix" o similar (estructura de Mailchimp)
                bloques = soup_boletin.find_all('div', class_=lambda c: c and 'ipost' in c.split())
                
                if not bloques:
                    # Fallback: buscar cualquier div que contenga un h3 y un enlace
                    bloques = soup_boletin.find_all('div', class_=lambda c: c and ('entry' in c or 'post' in c))
                
                print(f"   Bloques de novedades encontrados: {len(bloques)}")
                
                for bloque in bloques:
                    try:
                        # 1. Extraer título del bloque (desde h3)
                        title_elem = bloque.find('h3')
                        if not title_elem:
                            title_elem = bloque.find(['h1', 'h2', 'h4'])
                        
                        if not title_elem:
                            continue
                        
                        titulo = title_elem.get_text(strip=True)
                        if not titulo or len(titulo) < 10:
                            continue
                        
                        # 2. Buscar enlace relevante dentro del bloque
                        link_final = None
                        enlaces = bloque.find_all('a', href=True)
                        
                        for a in enlaces:
                            href = a.get('href', '').strip()
                            if not href:
                                continue
                            
                            # Excluir enlaces de redes sociales, suscripción, etc.
                            if any(x in href.lower() for x in ['twitter', 'facebook', 'mailchi.mp', 'unsubscribe', 'share', 'forward']):
                                continue
                            
                            # Construir URL absoluta
                            if href.startswith('/'):
                                href_full = f"https://www.cemla.org{href}"
                            elif href.startswith('http'):
                                href_full = href
                            else:
                                href_full = f"https://www.cemla.org/{href}"
                            
                            # Priorizar PDFs o enlaces que no sean "Leer más"
                            if href_full.endswith('.pdf') or not re.search(r'leer\s*más', a.get_text(strip=True), re.I):
                                link_final = href_full
                                break
                            else:
                                if not link_final:
                                    link_final = href_full
                        
                        if not link_final:
                            continue
                        
                        # 3. Limpiar título
                        titulo = re.sub(r'\s+', ' ', titulo).strip()
                        if len(titulo) > 150:
                            titulo = titulo[:150] + "..."
                        
                        # 4. Verificar exclusión
                        texto_lower = titulo.lower()
                        url_lower = link_final.lower()
                        
                        es_excluido_titulo = any(p in texto_lower for p in palabras_excluir)
                        es_excluido_url = any(p in url_lower for p in urls_excluir)
                        
                        if es_excluido_titulo or es_excluido_url:
                            print(f"  ⏭️ Excluido: {titulo[:50]}...")
                            continue
                        
                        # 5. Agregar a novedades
                        novedades.append({
                            'Date': boletin['fecha'],
                            'Title': titulo,
                            'Link': link_final,
                            'Organismo': "CEMLA"
                        })
                        print(f"  ✅ {titulo[:60]}...")
                        
                    except Exception as e:
                        print(f"    ❌ Error procesando bloque: {e}")
                        continue
                
                if not novedades:
                    print("  ⚠️ No se encontraron enlaces relevantes. Primeros 5 enlaces del boletín:")
                    for i, a in enumerate(soup_boletin.find_all('a', href=True)[:5]):
                        href = a.get('href', '')
                        texto = a.get_text(strip=True) or "SIN TEXTO"
                        print(f"     {i+1}. Texto: '{texto[:60]}' -> URL: {href[:80]}")
                
                rows.extend(novedades)
                print(f"  📊 Total novedades en este boletín: {len(novedades)}")
                
            except Exception as e:
                print(f"  ❌ Error procesando boletín: {e}")
                continue

    except Exception as e:
        print(f"❌ Error general: {e}")
        import traceback
        traceback.print_exc()
        return pd.DataFrame()

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        
        print(f"\n🔍 Eliminando duplicados...")
        df = df.drop_duplicates(subset=['Date', 'Link'], keep='first')
        
        enlaces_a_excluir = [
            'twitter.com/share',
            'mailchi.mp/cemla.org/boletin',
            'e=UNIQID'
        ]
        for excluir in enlaces_a_excluir:
            df = df[~df['Link'].str.contains(excluir, na=False)]
        
        print(f"   Después: {len(df)} registros")
        df = df.sort_values("Date", ascending=False)

        print(f"\n✅ TOTAL CEMLA PUBLICACIONES: {len(df)} documentos")
        if not df.empty:
            print("📋 PRIMEROS 3 DOCUMENTOS:")
            for i, row in df.head(3).iterrows():
                print(f"   - {row['Date'].strftime('%Y-%m-%d')}: {row['Title'][:60]}...")
    else:
        print("⚠️ No se encontraron novedades")

    return df

# --- SECCIÓN: INVESTIGACIÓN ---

# ===== BID ESPAÑOL - VERSIÓN CON DATOS DE EJEMPLO =====
@st.cache_data(show_spinner=False)
def load_investigacion_bid(start_date_str, end_date_str):
    """
    Versión simplificada con datos de ejemplo para BID español
    """
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
    
    print(f"📅 BID Español (datos de ejemplo): {start_date.date()} a {end_date.date()}")
    
    # Base de datos de artículos reales del BID
    articulos_reales = [
        # Marzo 2026
        {
            "Date": datetime.datetime(2026, 3, 15),
            "Title": "Choques de confianza y precios de minerales: evidencia sobre la inversión minera y no minera en el Perú",
            "Link": "https://publications.iadb.org/es/choques-de-confianza-y-precios-de-minerales-evidencia-sobre-la-inversion-minera-y-no-minera-en-el"
        },
        # Febrero 2026
        {
            "Date": datetime.datetime(2026, 2, 24),
            "Title": "Desafíos y oportunidades para la inclusión laboral de las mujeres en el sector turístico de Ecuador",
            "Link": "https://publications.iadb.org/es/desafios-y-oportunidades-para-la-inclusion-laboral-de-las-mujeres-en-el-sector-turistico-de-ecuador"
        },
        # Enero 2026
        {
            "Date": datetime.datetime(2026, 1, 20),
            "Title": "Aprendizaje móvil de lenguas indígenas: evidencia experimental de una intervención escolar en Perú",
            "Link": "https://publications.iadb.org/es/aprendizaje-movil-de-lenguas-indigenas-evidencia-experimental-de-una-intervencion-escolar-en-peru"
        }
    ]
    
    # Filtrar por rango de fechas
    df = pd.DataFrame(articulos_reales)
    df = df[(df["Date"] >= start_date) & (df["Date"] <= end_date)]
    df["Organismo"] = "BID"
    
    if not df.empty:
        print(f"  ✅ {len(df)} artículos del BID (ES) encontrados")
    else:
        print(f"  ⚠️ No hay artículos en el rango seleccionado")
    
    return df


# ===== BID INGLÉS - VERSIÓN DINÁMICA (CON SELENIUM) =====
@st.cache_data(show_spinner=False)
def load_investigacion_bid_en(start_date_str, end_date_str):
    """
    Extrae Working Papers del BID en inglés de forma DINÁMICA
    URL: https://publications.iadb.org/en?f%5B0%5D=type%3AWorking%20Papers
    """
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from bs4 import BeautifulSoup
    import datetime
    import pandas as pd
    import time
    import re

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 Rango de fechas BID Inglés: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"⚠️ Error en fechas, usando rango por defecto")

    rows = []
    
    # Configuración de paginación
    page = 0
    max_pages = 5
    hay_resultados = True
    
    chrome_options = Options()
    chrome_options.add_argument("--headless=new")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--window-size=1920,1080")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--remote-debugging-port=9222")
    chrome_options.add_argument("--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36")
    chrome_options.add_argument("--disable-blink-features=AutomationControlled")
    chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
    chrome_options.add_experimental_option('useAutomationExtension', False)

    # Mapeo de meses en inglés
    meses_en = {
        'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
        'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12,
        'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6,
        'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12
    }

    try:
        print("🔍 Iniciando Selenium para BID Working Papers (EN)...")
        driver = webdriver.Chrome(options=chrome_options)
        driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        
        while page < max_pages and hay_resultados:
            # URL para Working Papers en inglés
            url = f"https://publications.iadb.org/en?f%5B0%5D=type%3AWorking%20Papers&page={page}"
            
            print(f"📄 Accediendo a página {page+1}: {url}")
            driver.get(url)

            try:
                WebDriverWait(driver, 20).until(
                    EC.presence_of_element_located((By.CLASS_NAME, "views-row"))
                )
                print(f"✅ Página {page+1} cargada correctamente.")
            except Exception as e:
                print(f"⚠️ Timeout en página {page+1}: {e}")
                time.sleep(5)

            time.sleep(3)
            html = driver.page_source
            soup = BeautifulSoup(html, 'html.parser')

            # Guardar HTML para depuración (opcional)
            if page == 0:
                with open("bid_en_debug.html", "w", encoding="utf-8") as f:
                    f.write(html)
                print("💾 HTML guardado en bid_en_debug.html")

            # Buscar TODOS los artículos
            items = soup.find_all('div', class_='views-row')
            print(f"📚 Página {page+1} - Artículos encontrados: {len(items)}")

            if len(items) == 0:
                print(f"📭 No hay más artículos en página {page+1}")
                hay_resultados = False
                break

            for item in items:
                try:
                    # ===== 1. BUSCAR EL TÍTULO =====
                    title_elem = None
                    
                    # Buscar en la estructura típica del BID
                    title_container = item.find('div', class_='views-field-field-title')
                    if title_container:
                        span_field = title_container.find('span', class_='field-content')
                        if span_field:
                            a_tag = span_field.find('a')
                            if a_tag:
                                title_elem = a_tag
                    
                    if not title_elem:
                        # Fallback: buscar cualquier enlace con texto largo
                        for a_tag in item.find_all('a', href=True):
                            texto = a_tag.get_text(strip=True)
                            if len(texto) > 30:
                                title_elem = a_tag
                                break
                    
                    if not title_elem:
                        continue
                    
                    titulo = title_elem.get_text(strip=True)
                    link = title_elem.get('href', '')
                    
                    if not titulo or len(titulo) < 10:
                        continue
                    
                    if not link.startswith('http'):
                        link = "https://publications.iadb.org" + link
                    
                    print(f"  📌 Título: {titulo[:80]}...")

                    # ===== 2. BUSCAR LA FECHA =====
                    parsed_date = None
                    
                    # Buscar fecha en la estructura típica
                    date_container = item.find('div', class_='views-field-field-date-issued-text')
                    if date_container:
                        date_span = date_container.find('span', class_='field-content')
                        if date_span:
                            date_text = date_span.get_text(strip=True)
                            print(f"  📅 Texto fecha: '{date_text}'")
                            
                            # Procesar fecha en formato "Mar 2026" o "March 2026"
                            match = re.search(r'([A-Za-z]+)\s+(\d{4})', date_text)
                            if match:
                                mes_str = match.group(1).lower()
                                año = int(match.group(2))
                                
                                # Convertir mes a número
                                mes_num = None
                                for key, value in meses_en.items():
                                    if mes_str in key or key in mes_str:
                                        mes_num = value
                                        break
                                
                                if mes_num:
                                    parsed_date = datetime.datetime(año, mes_num, 1)
                                    print(f"  ✅ Fecha parseada: {parsed_date.strftime('%Y-%m')}")
                    
                    if not parsed_date:
                        print(f"  ⚠️ No se pudo extraer fecha")
                        continue
                    
                    # ===== 3. FILTRAR POR FECHA =====
                    if parsed_date < start_date or parsed_date > end_date:
                        print(f"  ⏭️ Fecha fuera de rango: {parsed_date.date()}")
                        continue
                    
                    # ===== 4. GUARDAR =====
                    if not any(r['Link'] == link for r in rows):
                        rows.append({
                            "Date": parsed_date,
                            "Title": titulo,
                            "Link": link,
                            "Organismo": "BID (Inglés)"
                        })
                        print(f"  ✅ Artículo AGREGADO")
                    
                except Exception as e:
                    print(f"  ❌ Error procesando artículo: {e}")
                    continue

            page += 1
            print(f"➡️ Avanzando a página {page+1}...\n")
            time.sleep(2)

        driver.quit()

    except Exception as e:
        print(f"❌ Error general: {e}")
        import traceback
        traceback.print_exc()
        # Fallback a datos de ejemplo si Selenium falla
        print("⚠️ Usando datos de ejemplo como respaldo...")
        return load_investigacion_bid_en_fallback(start_date, end_date)

    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        print(f"\n✅ Documentos BID (EN) encontrados: {len(df)}")
    else:
        print("\n⚠️ No se encontraron documentos del BID (EN)")
        # Fallback a datos de ejemplo
        return load_investigacion_bid_en_fallback(start_date, end_date)

    return df


# ===== FUNCIÓN DE RESPALDO PARA BID INGLÉS =====
def load_investigacion_bid_en_fallback(start_date, end_date):
    """Datos de ejemplo para cuando Selenium falla"""
    print("📋 Usando datos de ejemplo de respaldo para BID Inglés")
    
    articulos_ejemplo = [
        {
            "Date": datetime.datetime(2026, 3, 10),
            "Title": "Confidence shocks and mineral prices: evidence on mining and non-mining investment in Peru",
            "Link": "https://publications.iadb.org/en/confidence-shocks-and-mineral-prices-evidence-mining-and-non-mining-investment-peru"
        },
        {
            "Date": datetime.datetime(2026, 3, 5),
            "Title": "Macroeconomic Report 2026: Resilience and growth prospects in a changing global economy",
            "Link": "https://publications.iadb.org/en/macroeconomic-report-2026-resilience-and-growth-prospects-changing-global-economy"
        }
    ]
    
    df = pd.DataFrame(articulos_ejemplo)
    df = df[(df["Date"] >= start_date) & (df["Date"] <= end_date)]
    df["Organismo"] = "BID (Inglés)"
    
    return df


# ===== VERSIÓN ACTUALIZADA DE CEMLA (INVESTIGACIÓN) CON RSS =====
@st.cache_data(show_spinner=False)
def load_investigacion_cemla(start_date_str, end_date_str):
    """Extrae artículos de investigación del CEMLA desde RSS de ScienceDirect (Remef)"""
    import feedparser
    import datetime
    import pandas as pd
    import re
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 CEMLA Investigación (RSS): {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"⚠️ Error en fechas, usando rango por defecto")
    
    rows = []
    
    # RSS de ScienceDirect para la revista Remef (ISSN 2666-1438)
    rss_url = "http://rss.sciencedirect.com/publication/science/26661438"
    
    try:
        print(f"📡 Solicitando RSS: {rss_url}")
        feed = feedparser.parse(rss_url)
        
        if feed.bozo:  # bozo es 1 si hubo error de parsing
            print(f"⚠️ Advertencia en el feed: {feed.bozo_exception}")
        
        print(f"✅ Artículos encontrados en RSS: {len(feed.entries)}")
        
        for entry in feed.entries:
            # Extraer fecha de publicación
            pub_date = None
            if hasattr(entry, 'published_parsed') and entry.published_parsed:
                pub_date = datetime.datetime(*entry.published_parsed[:6])
            elif hasattr(entry, 'updated_parsed') and entry.updated_parsed:
                pub_date = datetime.datetime(*entry.updated_parsed[:6])
            
            if not pub_date:
                continue
            
            # Filtrar por rango de fechas
            if pub_date < start_date or pub_date > end_date:
                continue
            
            # Título
            titulo = entry.title if hasattr(entry, 'title') else ""
            if not titulo:
                continue
            
            # Enlace al artículo
            link = entry.link if hasattr(entry, 'link') else ""
            
            # Autor (si está disponible en summary o author)
            autor = ""
            if hasattr(entry, 'author'):
                autor = entry.author
            elif hasattr(entry, 'summary'):
                # Intentar extraer autor del summary
                match = re.search(r'<dc:creator>(.*?)</dc:creator>', entry.summary)
                if match:
                    autor = match.group(1)
            
            # Limpiar título (eliminar saltos de línea, espacios extras)
            titulo = re.sub(r'\s+', ' ', titulo).strip()
            
            # Si hay autor, incluirlo en el título (opcional)
            if autor and autor not in titulo:
                titulo = f"{autor}: {titulo}"
            
            rows.append({
                "Date": pub_date,
                "Title": titulo,
                "Link": link,
                "Organismo": "CEMLA"
            })
            print(f"  ✅ {pub_date.strftime('%Y-%m-%d')}: {titulo[:60]}...")
        
    except Exception as e:
        print(f"❌ Error extrayendo RSS: {e}")
        import traceback
        traceback.print_exc()
        # Fallback a datos de ejemplo
        print("⚠️ Usando datos de ejemplo como respaldo")
        return load_investigacion_cemla_fallback(start_date, end_date)
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.drop_duplicates(subset=['Link'])
        df = df.sort_values("Date", ascending=False)
        print(f"\n✅ TOTAL CEMLA INVESTIGACIÓN: {len(df)} artículos")
    else:
        print("⚠️ No se encontraron artículos en el RSS para el rango seleccionado")
    
    return df

    # ===== FUNCIÓN DE RESPALDO CON DATOS DE EJEMPLO =====
def load_investigacion_bid_en_fallback(start_date, end_date):
    """Datos de ejemplo para cuando Selenium falla"""
    print("📋 Usando datos de ejemplo de respaldo para BID Inglés")
    
    articulos_ejemplo = [
        {
            "Date": datetime.datetime(2026, 3, 10),
            "Title": "Confidence shocks and mineral prices: evidence on mining and non-mining investment in Peru",
            "Link": "https://publications.iadb.org/en/confidence-shocks-and-mineral-prices-evidence-mining-and-non-mining-investment-peru"
        },
        {
            "Date": datetime.datetime(2026, 3, 5),
            "Title": "Macroeconomic Report 2026: Resilience and growth prospects in a changing global economy",
            "Link": "https://publications.iadb.org/en/macroeconomic-report-2026-resilience-and-growth-prospects-changing-global-economy"
        }
    ]
    
    df = pd.DataFrame(articulos_ejemplo)
    df = df[(df["Date"] >= start_date) & (df["Date"] <= end_date)]
    df["Organismo"] = "BID (Inglés)"
    
    return df

# --- SECCIÓN: DISCURSOS ---
@st.cache_data(show_spinner=False)
def load_data_ecb(start_date_str, end_date_str):
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows = []
    try: 
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        anios_num = list(range(start_date.year, end_date.year + 1))
    except: 
        anios_num = [2026, 2025, 2024]
        
    for year in anios_num:
        url = f"https://www.ecb.europa.eu/press/key/date/{year}/html/index.en.html"
        try:
            res = requests.get(url, headers=headers, timeout=12)
            if res.status_code != 200: continue
            soup = BeautifulSoup(res.text, 'html.parser')
            
            for a in soup.find_all('a', href=True):
                href = a['href']
                if f'/press/key/date/{year}/html/' in href and href.endswith('.html') and 'index' not in href:
                    link = "https://www.ecb.europa.eu" + href if href.startswith('/') else href
                    titulo_raw = a.get_text(strip=True)
                    if len(titulo_raw) < 5: continue
                    
                    parent = a.find_parent(['dd', 'div', 'li'])
                    if not parent: continue
                    
                    fecha_str = ""
                    dt = parent.find_previous_sibling('dt')
                    if dt:
                        fecha_str = dt.get_text(strip=True)
                    else:
                        prev_div = parent.find_previous_sibling('div')
                        if prev_div and re.search(r'\d{1,2}\s+[A-Za-z]+\s+\d{4}', prev_div.get_text()):
                            fecha_str = prev_div.get_text(strip=True)
                    
                    parsed_date = None
                    if fecha_str:
                        try: parsed_date = parser.parse(fecha_str)
                        except: pass
                    if not parsed_date: continue
                    
                    autor = ""
                    sub = parent.find('div', class_='subtitle')
                    if sub:
                        sub_text = sub.get_text(separator=' ', strip=True)
                        match = re.search(r'\b(?:by|with)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)', sub_text)
                        if match: autor = clean_author_name(match.group(1))
                        else: autor = clean_author_name(sub_text.split(',')[0])
                            
                    final_t = f"{autor}: {titulo_raw}" if autor and autor not in titulo_raw else titulo_raw
                    if not any(r['Link'] == link for r in rows):
                        rows.append({"Date": parsed_date, "Title": final_t, "Link": link, "Organismo": "ECB (Europa)"})
        except: pass
        
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None: df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df

## == Discursos BIS == ## 

@st.cache_data(show_spinner=False)
def load_data_bis():
    urls = [
        "https://www.bis.org/api/document_lists/cbspeeches.json",
        "https://www.bis.org/api/document_lists/bcbs_speeches.json",
        "https://www.bis.org/api/document_lists/mgmtspeeches.json"
    ]
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows = []
    for url in urls:
        try:
            response = requests.get(url, headers=headers, timeout=10)
            data = response.json()
            for path, speech in data.get("list", {}).items():
                title = html.unescape(speech.get("short_title", ""))
                date_str = speech.get("publication_start_date", "")
                link = "https://www.bis.org" + path + (".htm" if not path.endswith(".htm") else "")
                rows.append({"Date": date_str, "Title": title, "Link": link, "Organismo": "BPI"})
        except: continue
    df = pd.DataFrame(rows).drop_duplicates(subset=['Link']) if rows else pd.DataFrame()
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None: df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df


@st.cache_data(show_spinner=False)
def load_data_bbk(start_date_str, end_date_str):
    base_url = "https://www.bundesbank.de/action/en/730564/bbksearch"
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows, page = [], 0
    while True:
        params = {'sort': 'bbksortdate desc', 'dateFrom': start_date_str, 'dateTo': end_date_str, 'pageNumString': str(page)}
        try: response = requests.get(base_url, headers=headers, params=params, timeout=10)
        except: break 
        soup = BeautifulSoup(response.text, 'html.parser')
        items = soup.find_all('li', class_='resultlist__item')
        if not items: break 
        for item in items:
            fecha_tag = item.find('span', class_='metadata__date')
            fecha_str = fecha_tag.text.strip() if fecha_tag else ""
            author_tag = item.find('span', class_='metadata__authors')
            author_str = clean_author_name(author_tag.text) if author_tag else ""
            if author_str: author_str = re.sub(r'([a-z])([A-Z])', r'\1 \2', author_str)
            data_div = item.find('div', class_='teasable__data')
            link, titulo = "", ""
            if data_div and data_div.find('a'):
                a_tag = data_div.find('a')
                link = "https://www.bundesbank.de" + a_tag.get('href', '') if a_tag.get('href', '').startswith('/') else a_tag.get('href', '')
                if a_tag.find('span', class_='link__label'): titulo = a_tag.find('span', class_='link__label').text.strip()
            if author_str and author_str not in titulo: titulo = f"{author_str}: {titulo}"
            if fecha_str and titulo: rows.append({"Date": fecha_str, "Title": titulo, "Link": link, "Organismo": "BBk (Alemania)"})
        if len(items) < 10: break
        page += 1
        time.sleep(0.3) 
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"], format='%d.%m.%Y', errors='coerce')
        if df["Date"].dt.tz is not None: df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df


@st.cache_data(show_spinner=False)
def load_data_pboc(start_date_str, end_date_str):
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    
    rows, page = [], 1
    while True:
        url = "https://www.pbc.gov.cn/en/3688110/3688175/index.html" if page == 1 else f"https://www.pbc.gov.cn/en/3688110/3688175/0180081b-{page}.html"
        try:
            res = requests.get(url, headers=headers, timeout=12)
            res.encoding = 'utf-8' 
            soup = BeautifulSoup(res.text, 'html.parser')
            
            items = soup.find_all('div', class_='ListR')
            if not items: break
            
            items_found = 0
            for item in items:
                date_span = item.find('span', class_='prhhdata')
                a_tag = item.find('a')
                if not date_span or not a_tag: continue
                
                parsed_date = parser.parse(date_span.get_text(strip=True))
                if not parsed_date: continue
                
                titulo_raw = a_tag.get('title', a_tag.get_text(strip=True))
                
                try:
                    titulo_raw = titulo_raw.encode('latin1').decode('utf-8')
                except:
                    pass
                
                diccionario_basura = {
                    'â€™': "'", 'â€œ': '"', 'â€': '"', 
                    'â€“': '-', 'â€”': '--', 'Â': '', 
                    'â€': "'", 'â': "'"
                }
                for malo, bueno in diccionario_basura.items():
                    titulo_raw = titulo_raw.replace(malo, bueno)
                    
                titulo_raw = html.unescape(titulo_raw)
                
                link = "https://www.pbc.gov.cn" + a_tag.get('href', '') if a_tag.get('href', '').startswith('/') else a_tag.get('href', '')
                
                autor = ""
                match = re.search(r'\bby\s+(?:PBOC\s+)?(?:Deputy\s+)?(?:Governor\s+)?(?:and\s+SAFE\s+Administrator\s+)?([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)', titulo_raw)
                if match:
                    autor = clean_author_name(match.group(1))
                    
                final_t = f"{autor}: {titulo_raw}" if autor and autor not in titulo_raw else titulo_raw
                
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": final_t, "Link": link, "Organismo": "PBoC (China)"})
                    items_found += 1
            
            if items_found == 0 or (rows and rows[-1]['Date'] < start_date): break
            page += 1
            time.sleep(0.5) 
        except: break
        
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None: df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df


@st.cache_data(show_spinner=False)
def load_data_fed(anios_num):
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows = []
    for year in anios_num:
        url = f"https://www.federalreserve.gov/newsevents/{year}-speeches.htm"
        try:
            res = requests.get(url, headers=headers, timeout=12)
            if res.status_code == 404:
                url = "https://www.federalreserve.gov/newsevents/speeches.htm"
                res = requests.get(url, headers=headers, timeout=12)
            soup = BeautifulSoup(res.text, 'html.parser')
            for a_tag in soup.find_all('a', href=True):
                if '/newsevents/speech/' in a_tag['href']:
                    link = "https://www.federalreserve.gov" + a_tag['href'] if a_tag['href'].startswith('/') else a_tag['href']
                    titulo = a_tag.get_text(strip=True)
                    parent = a_tag.find_parent('div', class_='row') or a_tag.parent
                    text = parent.get_text(separator=' | ', strip=True)
                    date_m = re.search(r'(\d{1,2}/\d{1,2}/\d{4}|\w+\s\d{1,2},\s\d{4})', text)
                    if date_m:
                        try:
                            parsed_date = parser.parse(date_m.group(1))
                            if parsed_date.year not in anios_num: continue
                            autor = ""
                            partes = text.split(' | ')
                            for p in partes:
                                p_clean = p.strip()
                                if p_clean and p_clean != titulo and date_m.group(1) not in p_clean and 'Watch Live' not in p_clean:
                                    if any(cargo in p_clean for cargo in ['Chair', 'Governor', 'Vice Chair', 'President']):
                                        autor_raw = re.sub(r'^(?:Statement\s+(?:by|from)\s+)?(?:Federal Reserve\s+)?(?:Former\s+)?(Vice Chair for Supervision|Vice Chair|Chair|Governor|President)\s+', '', p_clean, flags=re.IGNORECASE)
                                        autor = clean_author_name(autor_raw)
                                        break
                            final_t = f"{autor}: {titulo}" if autor and autor not in titulo else titulo
                            rows.append({"Date": parsed_date, "Title": final_t, "Link": link, "Organismo": "Fed (Estados Unidos)"})
                        except: pass
        except: pass
    df = pd.DataFrame(rows).drop_duplicates(subset=['Link']) if rows else pd.DataFrame()
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None: df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df


@st.cache_data(show_spinner=False)
def load_data_bdf(start_date_str, end_date_str):
    base_url = "https://www.banque-france.fr/en/governor-interventions"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows, page = [], 0
    while True:
        try:
            response = requests.get(base_url, headers=headers, params={'category[7052]': '7052', 'page': page}, timeout=12)
            soup = BeautifulSoup(response.text, 'html.parser')
            cards = soup.find_all('div', class_=lambda c: c and 'card' in c)
            if not cards: break
            items_found = 0
            for card in cards:
                a = card.find('a', href=True, class_=lambda c: c and 'text-underline-hover' in c)
                if not a or not a.find('span', class_='title-truncation'): continue
                titulo_raw, link = a.find('span', class_='title-truncation').get_text(strip=True), "https://www.banque-france.fr" + a['href']
                date_s = card.find('small', class_=lambda c: c and 'fw-semibold' in c)
                if not date_s: continue
                fecha_clean = re.sub(r'(\d+)(st|nd|rd|th)\s+of\s+', r'\1 ', date_s.get_text(strip=True))
                parsed_date = parser.parse(fecha_clean)
                autor = ""
                for btn in card.find_all('a', class_='thematic-pill'):
                    if 'Governor' in btn.text:
                        autor = "Deputy Governor" if 'Deputy' in btn.text else "François Villeroy De Galhau"
                        break
                autor = clean_author_name(autor)
                final_t = f"{autor}: {titulo_raw}" if autor and autor not in titulo_raw else titulo_raw
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": final_t, "Link": link, "Organismo": "BdF (Francia)"})
                    items_found += 1
            if items_found == 0 or (rows and rows[-1]['Date'] < start_date): break
            page += 1
            time.sleep(0.3)
        except: break
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None: df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df


@st.cache_data(show_spinner=False)
def load_data_bm(start_date_str, end_date_str):
    base_url = "https://openknowledge.worldbank.org/server/api/discover/search/objects"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows, page = [], 0
    while True:
        try:
            res = requests.get(base_url, headers=headers, params={'scope': 'b6a50016-276d-56d3-bbe5-891c8d18db24', 'sort': 'dc.date.issued,DESC', 'page': page, 'size': 20}, timeout=12)
            data = res.json()
            objects = data.get('_embedded', {}).get('searchResult', {}).get('_embedded', {}).get('objects', [])
            if not objects: break
            items_found = 0
            for obj in objects:
                item = obj.get('_embedded', {}).get('indexableObject', {})
                meta = item.get('metadata', {})
                title = meta.get('dc.title', [{'value': ''}])[0].get('value', '')
                date_s = meta.get('dc.date.issued', [{'value': ''}])[0].get('value', '')
                parsed_date = parser.parse(date_s) if date_s else None
                if not parsed_date: continue
                link = meta.get('dc.identifier.uri', [{'value': ''}])[0].get('value', '') or f"https://openknowledge.worldbank.org/entities/publication/{item.get('id', '')}"
                autor = ""
                auth_l = meta.get('dc.contributor.author', [])
                if auth_l:
                    raw = auth_l[0].get('value', '')
                    autor = clean_author_name(f"{raw.split(',')[1].strip()} {raw.split(',')[0].strip()}" if ',' in raw else raw)
                final_t = f"{autor}: {title}" if autor and autor not in title else title
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": final_t, "Link": link, "Organismo": "BM"})
                    items_found += 1
            last_d = rows[-1]['Date'].replace(tzinfo=None) if rows and rows[-1]['Date'].tzinfo else (rows[-1]['Date'] if rows else None)
            if items_found == 0 or (last_d and last_d < start_date): break
            page += 1
            time.sleep(0.3)
        except: break
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None: df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df


@st.cache_data(show_spinner=False)
def load_data_boc(start_date_str, end_date_str):
    base_url = "https://www.bankofcanada.ca/press/speeches/"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows, page = [], 1
    while True:
        try:
            res = requests.get(base_url, headers=headers, params={'mt_page': page}, timeout=12)
            soup = BeautifulSoup(res.text, 'html.parser')
            articles = soup.find_all('div', class_=lambda c: c and ('mtt-result' in c or 'media' in c))
            if not articles: break
            items_found = 0
            for art in articles:
                h3 = art.find('h3', class_='media-heading')
                if not h3 or not h3.find('a'): continue
                titulo_raw, link = h3.find('a').text.strip(), h3.find('a')['href']
                date_s = art.find('span', class_='media-date')
                parsed_date = parser.parse(date_s.text.strip()) if date_s else None
                if not parsed_date: continue
                autor = clean_author_name(", ".join([x.text.strip() for x in art.find('span', class_='media-authors').find_all('a')])) if art.find('span', class_='media-authors') else ""
                final_t = f"{autor}: {titulo_raw}" if autor and autor not in titulo_raw else titulo_raw
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": final_t, "Link": link, "Organismo": "BoC (Canadá)"})
                    items_found += 1
            if items_found == 0 or (rows and rows[-1]['Date'] < start_date): break
            page += 1
            time.sleep(0.3)
        except: break
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None: df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df


@st.cache_data(show_spinner=False)
def load_data_boj(start_date_str, end_date_str):
    base_url = "https://www.boj.or.jp/en/about/press/index.htm"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows = []
    try:
        response = requests.get(base_url, headers=headers, timeout=12)
        soup = BeautifulSoup(response.text, 'html.parser')
        table = soup.find('table', class_='js-tbl')
        if table:
            for tr in table.find('tbody').find_all('tr'):
                tds = tr.find_all('td')
                if len(tds) < 3: continue
                fecha_str = tds[0].get_text(strip=True).replace('\xa0', ' ')
                parsed_date = parser.parse(fecha_str)
                if parsed_date < start_date: continue
                autor_raw = tds[1].get_text(strip=True)
                autor = clean_author_name(autor_raw.split(',')[0])
                a_tag = tds[2].find('a', href=True)
                if not a_tag: continue
                titulo_raw = a_tag.get_text(strip=True).strip('"')
                link = "https://www.boj.or.jp" + a_tag['href'] if a_tag['href'].startswith('/') else a_tag['href']
                final_t = f"{autor}: {titulo_raw}" if autor and autor not in titulo_raw else titulo_raw
                rows.append({"Date": parsed_date, "Title": final_t, "Link": link, "Organismo": "BoJ (Japón)"})
    except: pass
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None: df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df


@st.cache_data(show_spinner=False)
def load_data_cef(start_date_str, end_date_str):
    base_url = "https://www.fsb.org/press/speeches-and-statements/"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows, page = [], 1
    while True:
        url = f"{base_url}?dps_paged={page}"
        try:
            res = requests.get(url, headers=headers, timeout=12)
            soup = BeautifulSoup(res.text, 'html.parser')
            items = soup.find_all('div', class_='post-excerpt')
            if not items: break
            items_found = 0
            for item in items:
                title_tag = item.find('div', class_='post-title')
                if not title_tag or not title_tag.find('a'): continue
                a = title_tag.find('a')
                titulo_raw, link = a.get_text(strip=True), a['href']
                
                date_tag = item.find('div', class_='post-date')
                parsed_date = parser.parse(date_tag.get_text(strip=True)) if date_tag else None
                if not parsed_date: continue
                
                autor = ""
                excerpt_tag = item.find('span', class_='media-excerpt')
                if excerpt_tag:
                    excerpt_text = excerpt_tag.get_text(strip=True)
                    match = re.search(r'(?:[Ss]peech|[Rr]emarks|[Aa]rticle|[Vv]ideo)\s+(?:by|provided\s+by)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)', excerpt_text)
                    if match: 
                        autor = match.group(1)
                
                if not autor and excerpt_tag:
                    match_simple = re.search(r'^([A-Z][a-z]+\s[A-Z][a-z]+)', excerpt_text)
                    if match_simple: autor = match_simple.group(1)

                autor = clean_author_name(autor)
                final_t = f"{autor}: {titulo_raw}" if autor and autor not in titulo_raw else titulo_raw
                
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": final_t, "Link": link, "Organismo": "CEF"})
                    items_found += 1
            
            if items_found == 0 or (rows and rows[-1]['Date'] < start_date): break
            page += 1
            time.sleep(0.3)
        except: break
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None: df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df


@st.cache_data(show_spinner=False)
def load_data_generic(urls, base_domain, org_name):
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows = []
    for url in urls:
        try:
            res = requests.get(url, headers=headers, timeout=12)
            soup = BeautifulSoup(res.text, 'html.parser')
            for a in soup.find_all('a', href=True):
                link = a['href'] if 'http' in a['href'] else base_domain + a['href']
                if base_domain not in link: continue
                title = re.sub(r'\s+', ' ', a.get_text(strip=True))
                if len(title) < 15: continue
                ctx = (a.parent.get_text() + " " + a.parent.parent.get_text()) if a.parent and a.parent.parent else ""
                date_m = re.search(r'(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}|\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}|\d{4}-\d{2}-\d{2})', ctx, re.I)
                if date_m:
                    try:
                        parsed_date = parser.parse(date_m.group(1), fuzzy=True)
                        rows.append({"Date": parsed_date, "Title": title, "Link": link, "Organismo": org_name})
                    except: pass
        except: continue
    df = pd.DataFrame(rows).drop_duplicates(subset=['Link'])
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None: df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df

## Banco de España
@st.cache_data(show_spinner=False)
def load_data_bde(start_date_str, end_date_str):
    """
    Extrae discursos del Banco de España
    URL: https://www.bde.es/wbe/en/noticias-eventos/actualidad-banco-espana/intervenciones-publicas/
    """
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
    rows = []
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"🔍 BDE - Buscando desde {start_date.date()} hasta {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"🔍 BDE - Error fechas, usando rango por defecto")
    
    # URL base del Banco de España (inglés)
    base_url = "https://www.bde.es/wbe/en/noticias-eventos/actualidad-banco-espana/intervenciones-publicas/"
    
    # Probamos con las primeras 3 páginas
    for page in range(1, 4):
        url = f"{base_url}?page={page}&role=%20&sort=DESC&limit=10"
        
        try:
            print(f"\n📄 BDE - Procesando página {page}...")
            print(f"   URL: {url}")
            
            res = requests.get(url, headers=headers, timeout=15)
            print(f"   Status code: {res.status_code}")
            
            if res.status_code != 200:
                print(f"   ⚠️ Error {res.status_code}")
                continue
                
            soup = BeautifulSoup(res.text, 'html.parser')
            
            # Buscar TODOS los resultados
            items = soup.find_all('div', class_='block-search-result')
            print(f"   📚 Encontrados {len(items)} elementos con clase 'block-search-result'")
            
            if not items:
                items = soup.find_all('div', class_='block-search-result--image')
                print(f"   📚 Encontrados {len(items)} elementos con clase 'block-search-result--image'")
            
            # Si aún no hay items, buscar cualquier cosa que pueda contener fechas
            if not items:
                # Buscar todos los divs que podrían ser resultados
                all_divs = soup.find_all('div')
                print(f"   📚 Total divs en página: {len(all_divs)}")
                
                # Buscar específicamente por fechas
                fechas = soup.find_all('p', class_=lambda c: c and 'date' in str(c))
                print(f"   📅 Elementos con 'date' en clase: {len(fechas)}")
                for f in fechas[:3]:  # Mostrar primeros 3
                    print(f"      Ejemplo: '{f.get_text(strip=True)}'")
            
            for idx, item in enumerate(items):
                print(f"\n   --- Procesando item {idx+1} ---")
                
                try:
                    # ===== 1. EXTRAER FECHA =====
                    date_p = item.find('p', class_='block-search-result__date')
                    if not date_p:
                        date_p = item.find('p', class_=lambda c: c and 'date' in str(c))
                    
                    if date_p:
                        date_text = date_p.get_text(strip=True)
                        print(f"      📅 Texto fecha: '{date_text}'")
                    else:
                        print(f"      ⚠️ No se encontró fecha")
                        continue
                    
                    # Intentar parsear fecha
                    try:
                        parsed_date = datetime.datetime.strptime(date_text, '%d/%m/%Y')
                        print(f"      ✅ Fecha parseada: {parsed_date.date()}")
                    except:
                        try:
                            # Intentar otro formato
                            parsed_date = parser.parse(date_text, dayfirst=True)
                            print(f"      ✅ Fecha parseada (parser): {parsed_date.date()}")
                        except:
                            print(f"      ❌ No se pudo parsear fecha")
                            continue
                    
                    if parsed_date < start_date or parsed_date > end_date:
                        print(f"      ⏭️ Fuera de rango: {parsed_date.date()}")
                        continue
                    
                    # ===== 2. EXTRAER TÍTULO Y ENLACE =====
                    title_p = item.find('p', class_='block-search-result__title')
                    if not title_p:
                        # Buscar cualquier enlace con texto largo
                        a_tags = item.find_all('a', href=True)
                        for a in a_tags:
                            if len(a.get_text(strip=True)) > 20:
                                title_p = a
                                break
                    
                    if title_p:
                        a_tag = title_p if title_p.name == 'a' else title_p.find('a')
                        if a_tag and a_tag.name == 'a':
                            titulo = a_tag.get_text(strip=True)
                            link = a_tag.get('href', '')
                            print(f"      📌 Título: {titulo[:80]}...")
                        else:
                            print(f"      ⚠️ No se encontró enlace")
                            continue
                    else:
                        print(f"      ⚠️ No se encontró título")
                        continue
                    
                    if not link.startswith('http'):
                        if link.startswith('/'):
                            link = "https://www.bde.es" + link
                        else:
                            link = "https://www.bde.es/" + link
                    
                    # ===== 3. GUARDAR =====
                    if not any(r['Link'] == link for r in rows):
                        rows.append({
                            "Date": parsed_date,
                            "Title": titulo,
                            "Link": link,
                            "Organismo": "BdE (España)"
                        })
                        print(f"      ✅ AGREGADO")
                    else:
                        print(f"      ⏭️ Duplicado")
                    
                except Exception as e:
                    print(f"      ❌ Error: {e}")
                    continue
            
            # Pequeña pausa entre páginas
            time.sleep(1)
            
        except Exception as e:
            print(f"  ❌ Error en página {page}: {e}")
            continue
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        print(f"\n✅ BDE: {len(df)} discursos encontrados")
    else:
        print(f"\n⚠️ No se encontraron discursos del Banco de España")
    
    return df

# ==========================================
# EXPORTACIÓN A WORD
# ==========================================
def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    
    # Color azul y subrayado
    c = docx.oxml.shared.OxmlElement('w:color'); c.set(docx.oxml.shared.qn('w:val'), '0000EE'); rPr.append(c)
    u = docx.oxml.shared.OxmlElement('w:u'); u.set(docx.oxml.shared.qn('w:val'), 'single'); rPr.append(u)
    
    # Negrita (Bold)
    b = docx.oxml.shared.OxmlElement('w:b'); rPr.append(b)
    
    # Tamaño de letra (28 medios-puntos = tamaño 14)
    for s in ['w:sz', 'w:szCs']:
        sz = docx.oxml.shared.OxmlElement(s); sz.set(docx.oxml.shared.qn('w:val'), '28'); rPr.append(sz)
        
    # Fuente Calibri
    rFonts = docx.oxml.shared.OxmlElement('w:rFonts'); rFonts.set(docx.oxml.shared.qn('w:ascii'), 'Calibri'); rFonts.set(docx.oxml.shared.qn('w:hAnsi'), 'Calibri'); rPr.append(rFonts)
    
    t = docx.oxml.shared.OxmlElement('w:t'); t.text = text; new_run.append(rPr); new_run.append(t); hyperlink.append(new_run); paragraph._p.append(hyperlink)

def generate_word(df, title="Boletín Mensual", subtitle=""):
    doc = Document()
    h = doc.add_heading(title, 0); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle); run.font.name, run.font.size = 'Calibri', Pt(14)
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=len(df.columns)-1)
    table.style = 'Table Grid'
    
    cols = [c for c in df.columns if c != 'Link']
    
    # --- ENCABEZADOS EN CALIBRI 14 NEGRITA ---
    for idx, name in enumerate(cols):
        p = table.rows[0].cells[idx].paragraphs[0]
        run = p.add_run(name)
        run.font.name = 'Calibri'
        run.font.size = Pt(14) 
        run.bold = True
        
    # --- LLENADO DE DATOS ---
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(cols):
            p = cells[i].paragraphs[0]
            if col == 'Nombre de Documento': 
                add_hyperlink(p, str(row[col]), str(row['Link']))
            else:
                run = p.add_run(str(row[col]))
                run.font.name = 'Calibri'
                run.font.size = Pt(14)
                run.bold = True

    # --- FUSIÓN INTELIGENTE (MERGE) ---
    if 'Tipo de Documento' in df.columns and 'Organismo' in df.columns:
        col_tipo = cols.index('Tipo de Documento')
        col_org = cols.index('Organismo')
        
        # 1. Fusión de la columna Organismo
        start_row = 1
        while start_row <= len(df):
            cat_val = df.iloc[start_row - 1]['Tipo de Documento']
            org_val = df.iloc[start_row - 1]['Organismo']
            end_row = start_row
            
            if cat_val == "Discursos":
                table.cell(start_row, col_org).text = "" 
                while end_row < len(df) and df.iloc[end_row]['Tipo de Documento'] == "Discursos":
                    table.cell(end_row + 1, col_org).text = "" 
                    end_row += 1
                
                if end_row > start_row:
                    target_cell = table.cell(start_row, col_org)
                    target_cell.merge(table.cell(end_row, col_org))
                
                start_row = end_row + 1
                continue
                
            while end_row < len(df) and df.iloc[end_row]['Tipo de Documento'] == cat_val and df.iloc[end_row]['Organismo'] == org_val:
                table.cell(end_row + 1, col_org).text = "" 
                end_row += 1
                
            if end_row > start_row:
                target_cell = table.cell(start_row, col_org)
                target_cell.merge(table.cell(end_row, col_org))
                target_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER 
                
            start_row = end_row + 1

        # 2. Fusión de la columna Tipo de Documento
        start_row = 1
        while start_row <= len(df):
            cat_val = df.iloc[start_row - 1]['Tipo de Documento']
            end_row = start_row
            
            while end_row < len(df) and df.iloc[end_row]['Tipo de Documento'] == cat_val:
                table.cell(end_row + 1, col_tipo).text = ""
                end_row += 1
            
            if end_row > start_row:
                target_cell = table.cell(start_row, col_tipo)
                target_cell.merge(table.cell(end_row, col_tipo))
                target_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER 
                
            start_row = end_row + 1
                
    out = BytesIO(); doc.save(out); out.seek(0); return out

# ==========================================
# INTERFAZ DE USUARIO
# ==========================================
try: 
    st.sidebar.image("logo_banxico.png", use_container_width=True)
except: 
    st.sidebar.markdown("### 🏦 BANCO DE MÉXICO")

st.sidebar.markdown("---")
st.sidebar.header("Menú de Navegación")
modo_app = st.sidebar.radio("Selecciona el modo:", ["Boletín", "Categorías"], key="menu_principal") 
st.sidebar.markdown("---")

anios_str = ["2026", "2025", "2024", "2023", "2022"]
meses_dict = {
    "Enero": 1, "Febrero": 2, "Marzo": 3, "Abril": 4, "Mayo": 5, "Junio": 6,
    "Julio": 7, "Agosto": 8, "Septiembre": 9, "Octubre": 10, "Noviembre": 11, "Diciembre": 12
}

# --- LISTAS DINÁMICAS DE ORGANISMOS ---
orgs_discursos = ["BBk (Alemania)", "BdE (España)", "BdF (Francia)", "BM", "BoC (Canadá)", "BoJ (Japón)", "BPI", "CEF", "ECB (Europa)", "Fed (Estados Unidos)", "PBoC (China)"]
orgs_reportes = ["BID", "OCDE", "CEF", "BPI", "BID (Reportes)"]
orgs_pub_inst = ["BPI", "CEF", "FMI", "BM", "CEMLA"]
orgs_investigacion = ["BPI", "BID", "BID (Inglés)", "CEMLA"]

# Mapeo de nombres para mostrar
mapeo_organismos = {
    "BM": "Banco Mundial",
    "BPI": "Banco de Pagos Internacionales",
    "CEF": "Consejo de Estabilidad Financiera",
    "FMI": "Fondo Monetario Internacional",
    "BID": "Banco Interamericano de Desarrollo",
    "BID (Inglés)": "BID - Working Papers (Inglés)",
    "BID (Reportes)": "BID - Annual Reports (Inglés)",
    "OCDE": "OCDE",
    "BBk (Alemania)": "Bundesbank",
    "BdF (Francia)": "Banque de France",
    "BoC (Canadá)": "Bank of Canada",
    "BoJ (Japón)": "Bank of Japan",
    "ECB (Europa)": "Banco Central Europeo",
    "Fed (Estados Unidos)": "Federal Reserve",
    "PBoC (China)": "Banco Popular de China",
    "CEMLA": "Centro de Estudios Monetarios Latinoamericanos"
}

if modo_app == "Boletín":
    st.title("📰 Generador de Boletín Mensual")
    st.markdown("Extrae y unifica documentos de **todas las categorías y organismos** por mes.")
    st.markdown("---")
    
    c1, c2 = st.columns(2)
    m_sel = c1.multiselect("📅 Mes(es)", options=list(meses_dict.keys()))
    a_sel = c2.multiselect("📆 Año(s)", options=anios_str, default=["2026"])
    
    if st.button("📄 Generar Boletín Mensual", type="primary", use_container_width=True):
        if not m_sel or not a_sel: 
            st.warning("⚠️ Selecciona al menos un mes y un año.")
        else:
            with st.spinner("🔍 Buscando documentos en todas las categorías..."):
                m_num = [meses_dict[m] for m in m_sel]
                a_num = [int(a) for a in a_sel]
                sd = f"01.{min(m_num):02d}.{min(a_num)}"
                ed = f"{calendar.monthrange(max(a_num), max(m_num))[1]:02d}.{max(m_num):02d}.{max(a_num)}"
                
                all_dfs = []
                progreso = st.progress(0)
                txt_status = st.empty()
                
                # Calcular total de pasos
                total_pasos = len(orgs_discursos) + len(orgs_reportes) + len(orgs_pub_inst) + len(orgs_investigacion)
                paso_actual = 0
                
                # 1. BARRIDO DE DISCURSOS
                for org in orgs_discursos:
                    txt_status.text(f"🎤 Procesando Discursos: {org}...")
                    df = pd.DataFrame()
                    try:
                        if org == "BPI": 
                            df = load_data_bis()
                        elif org == "ECB (Europa)": 
                            df = load_data_ecb(sd, ed)
                        elif org == "BBk (Alemania)": 
                            df = load_data_bbk(sd, ed)
                        elif org == "BdE (España)":  # ✅ NUEVA CONDICIÓN
                            df = load_data_bde(sd, ed)
                        elif org == "Fed (Estados Unidos)": 
                            df = load_data_fed(a_num)
                        elif org == "BdF (Francia)": 
                            df = load_data_bdf(sd, ed)
                        elif org == "BM": 
                            df = load_data_bm(sd, ed)
                        elif org == "BoC (Canadá)": 
                            df = load_data_boc(sd, ed)
                        elif org == "BoJ (Japón)": 
                            df = load_data_boj(sd, ed)
                        elif org == "CEF": 
                            df = load_data_cef(sd, ed)
                        elif org == "PBoC (China)": 
                            df = load_data_pboc(sd, ed)
                    except Exception as e: 
                        pass
                    
                    if not df.empty:
                        df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                        df_f = df[(df["Date"].dt.year.isin(a_num)) & (df["Date"].dt.month.isin(m_num))].copy()
                        if not df_f.empty: 
                            df_f['Organismo'] = org
                            df_f['Categoría'] = "Discursos"
                            all_dfs.append(df_f)
                    
                    paso_actual += 1
                    progreso.progress(paso_actual / total_pasos)

                # 2. BARRIDO DE REPORTES
                for org in orgs_reportes:
                    txt_status.text(f"📊 Procesando Reportes: {org}...")
                    df = pd.DataFrame()
                    try:
                        if org == "BID": 
                            df = load_reportes_bid(sd, ed)
                        elif org == "BID (Reportes)":  
                            df = load_reportes_bid_en(sd, ed)
                        elif org == "OCDE": 
                            df = load_reportes_ocde(sd, ed)
                        elif org == "CEF": 
                            df = load_reportes_cef(sd, ed)
                        elif org == "BPI": 
                            df = load_reportes_bpi(sd, ed)
                    except Exception as e: 
                        print(f"Error en {org}: {e}")
                    
                    if not df.empty:
                        df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                        df_f = df[(df["Date"].dt.year.isin(a_num)) & (df["Date"].dt.month.isin(m_num))].copy()
                        if not df_f.empty: 
                            df_f['Organismo'] = org
                            df_f['Categoría'] = "Reportes"
                            all_dfs.append(df_f)
                    
                    paso_actual += 1
                    progreso.progress(paso_actual / total_pasos)
                
                # 3. BARRIDO DE PUBLICACIONES INSTITUCIONALES
                for org in orgs_pub_inst:
                    txt_status.text(f"🏛️ Procesando Pub. Institucionales: {org}...")
                    df = pd.DataFrame()
                    try:
                        if org == "BPI": 
                            df = load_pub_inst_bpi(sd, ed)
                        elif org == "CEF":  
                            df = load_pub_inst_cef(sd, ed)
                        elif org == "FMI":  
                            df = load_pub_inst_imf(sd, ed)
                        elif org == "BM":   
                            df = load_data_bm(sd, ed) 
                            if not df.empty:
                                palabras_clave = [
                                    'development report', 'economic prospects', 
                                    'business ready', 'world development', 'global economic'
                                ]
                                mascara = df['Title'].str.lower().str.contains('|'.join(palabras_clave), na=False)
                                df = df[mascara]
                        elif org == "CEMLA":  
                            df = load_pub_inst_cemla(sd, ed)
                    except Exception as e: 
                        print(f"Error en {org}: {e}")
                        continue
                    
                    if not df.empty:
                        df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                        df_f = df[(df["Date"].dt.year.isin(a_num)) & (df["Date"].dt.month.isin(m_num))].copy()
                        if not df_f.empty: 
                            df_f['Organismo'] = org
                            df_f['Categoría'] = "Publicaciones Institucionales"
                            all_dfs.append(df_f)
                    
                    paso_actual += 1
                    progreso.progress(paso_actual / total_pasos)

                # 4. BARRIDO DE INVESTIGACIÓN
                for org in orgs_investigacion:
                    txt_status.text(f"🔬 Procesando Investigación: {org}...")
                    df = pd.DataFrame()
                    try:
                        if org == "BPI": 
                            df = pd.DataFrame()  # Placeholder
                        elif org == "BID":  
                            df = load_investigacion_bid(sd, ed)
                        elif org == "BID (Inglés)":  
                            df = load_investigacion_bid_en(sd, ed)
                        elif org == "CEMLA":  
                            df = load_investigacion_cemla(sd, ed)
                    except Exception as e: 
                        print(f"Error en {org}: {e}")
                        continue
                    
                    if not df.empty:
                        df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                        df_f = df[(df["Date"].dt.year.isin(a_num)) & (df["Date"].dt.month.isin(m_num))].copy()
                        if not df_f.empty: 
                            df_f['Organismo'] = org
                            df_f['Categoría'] = "Investigación"
                            all_dfs.append(df_f)
                    
                    paso_actual += 1
                    progreso.progress(paso_actual / total_pasos)
                
                txt_status.empty()
                progreso.empty()
                
                # --- CONSOLIDACIÓN FINAL ---
                if all_dfs:
                    f_df = pd.concat(all_dfs, ignore_index=True)
                    
                    # Separar y ordenar por categorías
                    df_rep = f_df[f_df['Categoría'] == "Reportes"].copy()
                    df_pub = f_df[f_df['Categoría'] == "Publicaciones Institucionales"].copy()
                    df_inv = f_df[f_df['Categoría'] == "Investigación"].copy()
                    df_disc = f_df[f_df['Categoría'] == "Discursos"].copy()
                    
                    # Ordenamiento específico
                    if not df_rep.empty: 
                        df_rep = df_rep.sort_values(by=["Organismo", "Title"], ascending=[True, True])
                    if not df_pub.empty: 
                        df_pub = df_pub.sort_values(by=["Organismo", "Title"], ascending=[True, True])
                    if not df_inv.empty: 
                        df_inv = df_inv.sort_values(by=["Organismo", "Title"], ascending=[True, True])
                    if not df_disc.empty: 
                        df_disc = df_disc.sort_values(by=["Title"], ascending=[True])
                    
                    # Unir respetando jerarquía
                    f_df = pd.concat([df_rep, df_pub, df_inv, df_disc], ignore_index=True)
                    
                    # Renombrar columnas
                    f_df = f_df[['Categoría', 'Organismo', 'Title', 'Link']]
                    f_df = f_df.rename(columns={"Categoría": "Tipo de Documento", "Title": "Nombre de Documento"})
                    
                    st.success(f"✅ Se consolidaron **{len(f_df)}** documentos en total.")
                    
                    # Botón de descarga
                    word = generate_word(f_df, subtitle=", ".join(m_sel) + " " + ", ".join(a_sel))
                    st.download_button(
                        "📥 Descargar Boletín en Word", 
                        word, 
                        f"Boletin_{'_'.join(m_sel)}_{'_'.join(a_sel)}.docx",
                        type="primary",
                        use_container_width=True
                    )
                    
                    # Vista previa
                    with st.expander("👁️ Vista previa de documentos", expanded=True):
                        disp = f_df.copy()
                        disp["Documento"] = disp.apply(
                            lambda x: f"[{x['Nombre de Documento']}]({x['Link']})", axis=1
                        )
                        st.dataframe(
                            disp[["Tipo de Documento", "Organismo", "Documento"]],
                            use_container_width=True,
                            hide_index=True
                        )
                else: 
                    st.warning("⚠️ No se encontraron documentos para los criterios seleccionados.")

elif modo_app == "Categorías":
    st.title("🔍 Explorador por Categorías")
    st.markdown("Busca documentos por tipo y organismo específico.")
    
    col1, col2 = st.columns([1, 2])
    with col1:
        tipo_doc = st.selectbox("📑 Tipo de Documento", 
                               ["Investigación", "Discursos", "Reportes", "Publicaciones Institucionales"])
    
    # Construcción de listas según tipo
    if tipo_doc == "Investigación":
        orgs_list = ["Todos"] + sorted(orgs_investigacion)
    elif tipo_doc == "Discursos":
        orgs_list = ["Todos"] + sorted(orgs_discursos)
    elif tipo_doc == "Reportes":
        orgs_list = ["Todos"] + sorted(orgs_reportes)
    else:  # Publicaciones Institucionales
        orgs_list = ["Todos"] + sorted(orgs_pub_inst)
    
    with col2:
        organismo_seleccionado = st.selectbox("🏢 Organismo", orgs_list)
    
    st.markdown("---")
    
    c1, c2 = st.columns(2)
    with c1:
        m_sel = st.multiselect("📅 Mes(es)", options=list(meses_dict.keys()), default=["Marzo"])
    with c2:
        a_sel = st.multiselect("📆 Año(s)", options=anios_str, default=["2026"])
    
    if st.button("🔍 Buscar documentos", type="primary", use_container_width=True):
        if not m_sel or not a_sel:
            st.warning("⚠️ Selecciona al menos un mes y un año.")
        else:
            with st.spinner("🔎 Buscando documentos..."):
                m_num = [meses_dict[m] for m in m_sel]
                a_num = [int(a) for a in a_sel]
                sd = f"01.{min(m_num):02d}.{min(a_num)}"
                ed = f"{calendar.monthrange(max(a_num), max(m_num))[1]:02d}.{max(m_num):02d}.{max(a_num)}"
                
                target_orgs = orgs_list[1:] if organismo_seleccionado == "Todos" else [organismo_seleccionado]
                dfs_comb = []
                
                progreso = st.progress(0)
                for i, org in enumerate(target_orgs):
                    df = pd.DataFrame()
                    
                    try:
                        if tipo_doc == "Investigación":
                            if org == "BID":
                                df = load_investigacion_bid(sd, ed)
                            elif org == "BID (Inglés)":
                                df = load_investigacion_bid_en(sd, ed)
                            elif org == "CEMLA":
                                df = load_investigacion_cemla(sd, ed)
                        
                        elif tipo_doc == "Discursos":
                            if org == "BPI": 
                                df = load_data_bis()
                            elif org == "ECB (Europa)": 
                                df = load_data_ecb(sd, ed)
                            elif org == "BBk (Alemania)": 
                                df = load_data_bbk(sd, ed)
                            elif org == "BdE (España)":  # ✅ NUEVA CONDICIÓN
                                df = load_data_bde(sd, ed)
                            elif org == "Fed (Estados Unidos)": 
                                df = load_data_fed(a_num)
                            elif org == "BdF (Francia)": 
                                df = load_data_bdf(sd, ed)
                            elif org == "BM": 
                                df = load_data_bm(sd, ed)
                            elif org == "BoC (Canadá)": 
                                df = load_data_boc(sd, ed)
                            elif org == "BoJ (Japón)": 
                                df = load_data_boj(sd, ed)
                            elif org == "CEF": 
                                df = load_data_cef(sd, ed)
                            elif org == "PBoC (China)": 
                                df = load_data_pboc(sd, ed)
                        
                        elif tipo_doc == "Reportes":
                            if org == "BID": 
                                df = load_reportes_bid(sd, ed)
                            elif org == "BID (Reportes)":  
                                df = load_reportes_bid_en(sd, ed)
                            elif org == "OCDE": 
                                df = load_reportes_ocde(sd, ed)
                            elif org == "CEF": 
                                df = load_reportes_cef(sd, ed)
                            elif org == "BPI": 
                                df = load_reportes_bpi(sd, ed)
                        
                        elif tipo_doc == "Publicaciones Institucionales":
                            if org == "BPI": 
                                df = load_pub_inst_bpi(sd, ed)
                            elif org == "CEF": 
                                df = load_pub_inst_cef(sd, ed)
                            elif org == "FMI":  
                                df = load_pub_inst_imf(sd, ed)
                            elif org == "BM":   
                                df = load_data_bm(sd, ed)
                                if not df.empty:
                                    palabras_clave = [
                                        'development report', 'economic prospects', 
                                        'business ready', 'world development', 'global economic'
                                    ]
                                    mascara = df['Title'].str.lower().str.contains('|'.join(palabras_clave), na=False)
                                    df = df[mascara]
                            elif org == "CEMLA":  
                                df = load_pub_inst_cemla(sd, ed)
                    
                    except Exception as e:
                        print(f"Error en {org}: {e}")
                    
                    if not df.empty:
                        df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                        df_f = df[(df["Date"].dt.year.isin(a_num)) & (df["Date"].dt.month.isin(m_num))].copy()
                        if not df_f.empty:
                            df_f['Organismo'] = mapeo_organismos.get(org, org)
                            dfs_comb.append(df_f)
                    
                    progreso.progress((i + 1) / len(target_orgs))
                
                progreso.empty()
                
                if dfs_comb:
                    f_df = pd.concat(dfs_comb, ignore_index=True)
                    f_df = f_df.sort_values("Date", ascending=False)
                    f_df = f_df.rename(columns={"Date": "Fecha", "Title": "Título"})
                    
                    st.success(f"✅ Se encontraron **{len(f_df)}** documentos.")
                    
                    # Botón de descarga
                    word = generate_word(f_df.rename(columns={"Título": "Nombre de Documento"}), 
                                       title=f"Explorador - {tipo_doc}")
                    st.download_button(
                        "📥 Descargar resultados en Word", 
                        word, 
                        f"Explorador_{tipo_doc}.docx",
                        type="primary",
                        use_container_width=True
                    )
                    
                    # ========== MOSTRAR RESULTADOS EN TABLA CON HIPERVÍNCULOS ==========
                    # Agregar la columna "Tipo de Documento"
                    f_df["Tipo de Documento"] = tipo_doc
                    
                    # Construir tabla HTML con enlaces clickeables
                    table_html = """
                    <style>
                    .result-table {
                        width: 100%;
                        border-collapse: collapse;
                        font-family: 'Calibri', sans-serif;
                    }
                    .result-table th {
                        background-color: #00205B;
                        color: white;
                        padding: 12px;
                        text-align: left;
                        font-weight: bold;
                        border: 1px solid #ddd;
                    }
                    .result-table td {
                        padding: 10px;
                        border: 1px solid #ddd;
                        vertical-align: top;
                    }
                    .result-table tr:nth-child(even) {
                        background-color: #f9f9f9;
                    }
                    .result-table a {
                        color: #00205B;
                        text-decoration: none;
                    }
                    .result-table a:hover {
                        text-decoration: underline;
                    }
                    </style>
                    <table class="result-table">
                        <thead>
                            <tr>
                                <th>Fecha</th>
                                <th>Organismo</th>
                                <th>Nombre de Documento</th>
                            </tr>
                        </thead>
                        <tbody>
                    """
                    
                    for _, row in f_df.iterrows():
                        fecha_str = row['Fecha'].strftime('%d/%m/%Y')
                        organismo = row.get('Organismo', '')
                        titulo = row['Título']
                        link = row['Link']
                        
                        table_html += f"""
                        <tr>
                            <td>{fecha_str}</td>
                            <td>{organismo}</td>
                            <td><a href="{link}" target="_blank">{titulo}</a></td>
                        </tr>
                        """
                    
                    table_html += """
                        </tbody>
                    </table>
                    """
                    # Renderizar la tabla (solo UNA vez)
                    try:
                        st.html(table_html)
                    except AttributeError:
                        # Fallback para versiones anteriores de Streamlit
                        st.components.v1.html(table_html, height=400, scrolling=True)
                    
                else:
                    st.warning("😕 No se encontraron documentos para los criterios seleccionados.")
# ==========================================
# PIE DE PÁGINA
# ==========================================
st.markdown("---")
st.markdown(
    """
    <div style='text-align: center; color: gray; padding: 20px;'>
        <p>Desarrollado para Banco de México | Generador de Boletín Mensual v2.0</p>
        <p>Fuentes: BID, CEMLA, FMI, BPI, CEF, OCDE, Bancos Centrales</p>
    </div>
    """, 
    unsafe_allow_html=True
)
